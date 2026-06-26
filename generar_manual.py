import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globales ─────────────────────────────────────────────────────────
sections = doc.sections
section = sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2)

NEGRO    = RGBColor(0x1a, 0x1a, 0x2e)
ROJO     = RGBColor(0xc0, 0x39, 0x2b)
GRIS     = RGBColor(0x55, 0x56, 0x57)
GRIS_CLR = RGBColor(0xf2, 0xf2, 0xf2)
AZUL     = RGBColor(0x1a, 0x73, 0xe8)

def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    val.get('sz',  '6'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=16, color=ROJO)
        # línea inferior roja
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'C0392B')
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        set_font(run, bold=True, size=13, color=NEGRO)
    else:
        set_font(run, bold=True, size=11, color=GRIS)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_steps(doc, steps):
    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.6)
        r1 = p.add_run(f'{title}: ' if title else '')
        set_font(r1, bold=True, size=11)
        r2 = p.add_run(desc)
        set_font(r2, size=11)

def add_tip(doc, text, warn=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.6)
    icon = '⚠ ' if warn else 'ℹ '
    col  = RGBColor(0x85, 0x64, 0x04) if warn else RGBColor(0x0d, 0x47, 0xa1)
    r1 = p.add_run(icon)
    set_font(r1, bold=True, size=10, color=col)
    r2 = p.add_run(text)
    set_font(r2, italic=True, size=10, color=col)

def add_flow(doc, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.6)
    for j, item in enumerate(items):
        r = p.add_run(item)
        set_font(r, bold=True, size=10, color=ROJO if '→' not in item else GRIS)
        if j < len(items) - 1 and '→' not in item:
            r2 = p.add_run(' → ')
            set_font(r2, size=10, color=GRIS)

# ════════════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('MANUAL DE PROCEDIMIENTO')
set_font(r, bold=True, size=26, color=ROJO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sistema de Costos Gastronómicos')
set_font(r, bold=True, size=18, color=NEGRO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Instituto Culinary')
set_font(r, size=14, color=GRIS)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run(f'Versión 1.0  ·  {datetime.date.today().strftime("%B %Y").capitalize()}')
set_font(r, size=11, color=GRIS)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ÍNDICE (manual — Word no genera TOC automático sin macros)
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'ÍNDICE DE CONTENIDOS', 1)

TOC = [
    ('1.', 'Introducción y flujo de datos'),
    ('2.', 'Acceso al sistema y roles de usuario'),
    ('3.', 'Datos base'),
    ('   3.1', 'Ingredientes'),
    ('   3.2', 'Proveedores'),
    ('   3.3', 'Tipos de Taller'),
    ('4.', 'Operaciones diarias'),
    ('   4.1', 'Sesiones de Clase'),
    ('   4.2', 'Inventario'),
    ('   4.3', 'Lista de Compras'),
    ('5.', 'Planificación académica'),
    ('   5.1', 'Semanas'),
    ('   5.2', 'Módulos Académicos'),
    ('   5.3', 'Carreras / Especialidades'),
    ('6.', 'Análisis y reportes'),
    ('   6.1', 'Dashboard (KPIs)'),
    ('   6.2', 'Simulador de Costos'),
    ('   6.3', 'Comparador de Precios'),
    ('   6.4', 'Historial de Precios'),
    ('   6.5', 'Presupuesto vs Real'),
    ('   6.6', 'Resumen por Docente'),
    ('   6.7', 'Comparativa de Períodos'),
    ('7.', 'Documentos'),
    ('   7.1', 'Orden de Compra'),
    ('   7.2', 'Informe PDF por Carrera'),
    ('8.', 'Herramientas'),
    ('   8.1', 'Búsqueda Global'),
    ('   8.2', 'Importar desde Excel'),
    ('   8.3', 'Backup / Exportar datos'),
    ('9.', 'Configuración y accesibilidad'),
]
for num, title in TOC:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    lv  = num.startswith('   ')
    p.paragraph_format.left_indent = Cm(1.2) if lv else Cm(0)
    r1  = p.add_run(num.strip() + '  ')
    set_font(r1, bold=not lv, size=11, color=ROJO if not lv else NEGRO)
    r2  = p.add_run(title)
    set_font(r2, size=11, color=NEGRO if not lv else GRIS)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN Y FLUJO DE DATOS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. INTRODUCCIÓN Y FLUJO DE DATOS', 1)
add_body(doc,
    'El Sistema de Costos Gastronómicos de Instituto Culinary es una aplicación web '
    'que permite registrar, calcular y analizar los costos de materia prima de cada '
    'clase práctica. Los costos se propagan automáticamente desde el ingrediente '
    'individual hasta el costo total por alumno de una carrera completa.')

add_heading(doc, 'Cadena de cálculo automático', 2)
add_flow(doc, ['Ingrediente', 'Sesión de Clase', 'Taller (tipo)', 'Módulo', 'Carrera'])
add_body(doc,
    'Cada eslabón recibe el costo del anterior y lo agrega a su nivel. '
    'El resultado final es el costo de materia prima por alumno en cada carrera.')

add_heading(doc, 'Orden de trabajo recomendado (primer uso)', 2)
add_steps(doc, [
    ('Categorías y proveedores', 'son la base para clasificar ingredientes. Créalos primero.'),
    ('Ingredientes',             'carga cada insumo con precio neto, unidad y proveedor.'),
    ('Tipos de taller',          'define los tipos de clases (Cocina Caliente, Panadería, etc.).'),
    ('Semanas',                  'crea los períodos lectivos del año académico.'),
    ('Sesiones de clase',        'registra cada clase con sus ingredientes y cantidades usadas.'),
    ('Inventario',               'ingresa el stock real mensualmente.'),
    ('Módulos académicos',       'agrupa tipos de taller con número de sesiones planificadas.'),
    ('Carreras',                 'vincula módulos y configura el overhead % para obtener el costo final.'),
])
add_tip(doc, 'Una vez cargados los datos, el Dashboard y todos los análisis se actualizan en tiempo real sin pasos adicionales.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. ACCESO Y ROLES
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. ACCESO AL SISTEMA Y ROLES DE USUARIO', 1)

add_heading(doc, 'Iniciar sesión', 2)
add_steps(doc, [
    ('Acceso', 'Ingresa a https://culinary.pythonanywhere.com/ desde cualquier navegador.'),
    ('Credenciales', 'Escribe tu usuario y contraseña en la pantalla de inicio.'),
    ('Autenticación', 'El sistema verifica las credenciales contra la base de datos con cifrado SHA-256. Si son correctas, la sesión se mantiene activa en el navegador.'),
    ('Persistencia', 'No necesitas volver a iniciar sesión al recargar la página. Para cerrar sesión presiona el botón ⏏ en el pie del menú lateral.'),
])

add_heading(doc, 'Roles disponibles', 2)

tbl = doc.add_table(rows=3, cols=3)
tbl.style = 'Table Grid'
headers = ['Rol', 'Permisos', 'Perfil recomendado']
for j, h in enumerate(headers):
    c = tbl.rows[0].cells[j]
    shade_cell(c, '1A1A2E')
    p2 = c.paragraphs[0]
    r  = p2.add_run(h)
    set_font(r, bold=True, size=10, color=RGBColor(0xff, 0xff, 0xff))

rows_data = [
    ('👑 Admin',    'Acceso total: crear, editar, eliminar, gestionar usuarios, importar datos.',
                    'Jefe de área / administrador del sistema.'),
    ('👁 Lectura',  'Solo visualización. Botones de acción ocultos. No puede crear ni modificar.',
                    'Docentes y personal de apoyo.'),
]
for i, (rol, perm, perf) in enumerate(rows_data, 1):
    for j, text in enumerate([rol, perm, perf]):
        c = tbl.rows[i].cells[j]
        if j == 0:
            shade_cell(c, 'FEF9C3')
        p2 = c.paragraphs[0]
        r  = p2.add_run(text)
        set_font(r, size=10, bold=(j == 0))

add_tip(doc, 'Los docentes deben tener rol Lectura. Solo el administrador debe tener rol Admin.')

add_heading(doc, 'Gestión de usuarios', 2)
add_steps(doc, [
    ('Acceder',   'Ve a Admin → Usuarios en el menú lateral (visible solo para admins).'),
    ('Crear',     'Haz clic en + Nuevo usuario. Ingresa nombre de usuario (minúsculas), nombre completo, contraseña (mín. 6 caracteres) y rol.'),
    ('Editar',    'Haz clic en el ícono ✏ del usuario. Si dejas la contraseña en blanco, no se modifica.'),
    ('Desactivar','Edita el usuario y cambia el estado a Inactivo. No se elimina, solo se bloquea el acceso.'),
])
add_tip(doc, 'El usuario "admin" no puede eliminarse para garantizar acceso de recuperación.', warn=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. DATOS BASE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. DATOS BASE', 1)

add_heading(doc, '3.1 Ingredientes', 2)
add_body(doc,
    'Maestro de materias primas. Todo costo del sistema se calcula a partir de los '
    'precios registrados aquí. Mantenerlos actualizados es crítico para la exactitud '
    'de los análisis.')
add_steps(doc, [
    ('Crear',         'Haz clic en + Nuevo ingrediente.'),
    ('Campos básicos','Ingresa: código interno, nombre (en mayúsculas recomendado), precio neto en CLP, unidad de medida (kg, L, UN, etc.).'),
    ('Clasificación', 'Asigna categoría y proveedor (deben existir previamente).'),
    ('Stock Mínimo',  'Define el umbral de alerta. Si el inventario cae por debajo de este valor, aparecerá un punto rojo en la tabla.'),
    ('Filtros',       'Usa el selector de categoría para navegar listados extensos.'),
    ('Actualización', 'Edita el ingrediente (ícono ✏) y cambia el precio. La fecha de última actualización queda registrada automáticamente.'),
])
add_tip(doc, 'Actualiza los precios al menos una vez por mes. Un precio desactualizado distorsiona todos los cálculos de sesiones y carreras.')

add_heading(doc, '3.2 Proveedores', 2)
add_body(doc,
    'Directorio de proveedores de insumos. Los datos aquí registrados se usan en las '
    'órdenes de compra automáticas.')
add_steps(doc, [
    ('Crear',        'Haz clic en + Nuevo proveedor.'),
    ('Datos',        'Ingresa: número de proveedor, nombre, tipo (Mayorista / Minorista), sede y contacto.'),
    ('Condición pago','Define la condición de pago: 30 días, 60 días, contado, etc.'),
    ('Contacto',     'Agrega correo y teléfono — se incluyen en las órdenes de compra impresas.'),
    ('Desactivar',   'Un proveedor inactivo no aparece en nuevos registros pero su historial se conserva.'),
])

add_heading(doc, '3.3 Tipos de Taller', 2)
add_body(doc,
    'Categorías de clases prácticas. Cada sesión pertenece a un tipo de taller, '
    'y los costos se agrupan por este criterio en los análisis.')
add_steps(doc, [
    ('Crear',  'Haz clic en + Nuevo taller.'),
    ('Código', 'Asigna un código corto único (ej: CC01, PA01).'),
    ('Color',  'Elige un color identificador — aparecerá como badge en sesiones y gráficos.'),
])
add_tip(doc, 'Ejemplos típicos: Cocina Caliente, Cocina Fría, Panadería y Repostería, Cocina Internacional, Bar y Coctelería.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. OPERACIONES DIARIAS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. OPERACIONES DIARIAS', 1)

add_heading(doc, '4.1 Sesiones de Clase', 2)
add_body(doc,
    'Registro central del sistema. Cada entrada representa una clase ejecutada '
    'con sus ingredientes y cantidades reales. Es el origen de todos los costos calculados.')
add_steps(doc, [
    ('Nueva sesión',      'Haz clic en + Nueva sesión.'),
    ('Datos de la clase', 'Ingresa: código de clase, tipo de taller, semana, fecha y número de alumnos presentes.'),
    ('Profesor',          'Registra el nombre exacto del docente que dictó la clase (se usa en el Resumen por Docente).'),
    ('Agregar ingredientes','Haz clic en + Ingrediente. Por cada uno ingresa:\n   • Cant/Alumno: cantidad por persona (el total se calcula automáticamente).\n   • Total: puedes editarlo manualmente si la cantidad real difiere.'),
    ('Guardar',           'El costo total y costo/alumno se calculan automáticamente por el servidor.'),
    ('Ver desglose',      'Haz clic en el ícono 👁 de cualquier sesión para ver el detalle completo de ingredientes y costos.'),
])
add_tip(doc, 'Usa nombres de docente consistentes (ej: siempre "Ana García", nunca "A. García" o "ana garcia") para que el Resumen por Docente agrupe correctamente.')

add_heading(doc, '4.2 Inventario', 2)
add_body(doc,
    'Registro del stock real de cada ingrediente. Se actualiza manualmente cada mes '
    'y sirve de base para la Lista de Compras y las alertas de stock mínimo.')
add_steps(doc, [
    ('Actualizar stock', 'Haz clic en + Actualizar stock.'),
    ('Seleccionar',      'Elige el ingrediente en el desplegable.'),
    ('Cantidad',         'Ingresa la cantidad actual en la unidad de medida del ingrediente.'),
    ('Período',          'Selecciona mes y año del conteo físico.'),
    ('Guardar',          'El sistema compara el stock con la demanda calculada de sesiones para generar la lista de compras.'),
])
add_tip(doc, 'Actualizar el inventario mensualmente es necesario para que la Lista de Compras y las alertas de stock mínimo funcionen correctamente.', warn=True)

add_heading(doc, '4.3 Lista de Compras', 2)
add_body(doc,
    'Generada automáticamente comparando la demanda de sesiones futuras contra el '
    'stock disponible en inventario. No requiere configuración adicional.')
add_steps(doc, [
    ('Acceder',    'Navega a Lista de Compras en el menú lateral.'),
    ('Leer tabla', 'Columnas clave: Necesario (demanda proyectada), Stock (inventario actual), A Comprar (diferencia), Costo Est. (proyección de gasto).'),
    ('Urgentes',   'Los ítems en rojo con estado COMPRAR son los más críticos.'),
    ('Exportar',   'Haz clic en 📊 Excel para descargar la lista completa.'),
    ('Orden formal','Para generar una orden de compra agrupada por proveedor, ve al módulo Orden de Compra.'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. PLANIFICACIÓN ACADÉMICA
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. PLANIFICACIÓN ACADÉMICA', 1)

add_heading(doc, '5.1 Semanas', 2)
add_body(doc, 'Períodos semanales del año académico. Organizan las sesiones temporalmente y permiten filtrar por período.')
add_steps(doc, [
    ('Crear',       'Haz clic en + Nueva semana.'),
    ('Número',      'Ingresa el número de semana (1 a 53) y el año académico.'),
    ('Nombre',      'Agrega una descripción opcional (ej: "Semana 1 – Inducción").'),
    ('Fechas',      'Define fechas de inicio y fin para tener referencia temporal precisa.'),
    ('Uso',         'Las semanas creadas aparecen en el selector al registrar nuevas sesiones y en los filtros del dashboard.'),
])

add_heading(doc, '5.2 Módulos Académicos', 2)
add_body(doc,
    'Un módulo agrupa tipos de talleres bajo una unidad curricular. Es el eslabón '
    'intermedio entre el taller y la carrera en la cadena de costos. Un módulo puede '
    'pertenecer a múltiples carreras.')
add_steps(doc, [
    ('Crear',     'Haz clic en + Nuevo módulo.'),
    ('Datos',     'Ingresa nombre, descripción y horas totales del módulo.'),
    ('Talleres',  'Haz clic en + Agregar taller. Por cada taller indica el número de sesiones planificadas en el módulo.'),
    ('Reutilizar','El mismo módulo puede asignarse a varias carreras sin duplicarlo.'),
])
add_tip(doc, 'Ejemplo: el módulo "Fundamentos Culinarios" puede tener 8 sesiones de Cocina Caliente + 4 sesiones de Cocina Fría.')

add_heading(doc, '5.3 Carreras / Especialidades', 2)
add_body(doc,
    'Programa académico completo. El costo total/alumno es el resultado final de toda '
    'la cadena de cálculo y el dato más relevante para la dirección.')
add_steps(doc, [
    ('Crear',    'Haz clic en + Nueva carrera.'),
    ('Datos',    'Ingresa nombre y duración total en horas.'),
    ('Overhead', 'Define el porcentaje de costos indirectos (ej: 20%). El sistema lo aplica sobre el costo de materia prima.'),
    ('Resultado','El sistema calcula automáticamente:\n   • Costo Materia Prima: suma módulos → talleres → sesiones → ingredientes.\n   • Costo Total/Alumno: materia prima + overhead %.'),
    ('Informe',  'Para generar el informe imprimible ve a Backup → Informe PDF por Carrera.'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. ANÁLISIS Y REPORTES
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. ANÁLISIS Y REPORTES', 1)

add_heading(doc, '6.1 Dashboard (KPIs)', 2)
add_body(doc,
    'Pantalla principal del sistema. Muestra los indicadores clave y gráficos de resumen '
    'calculados en tiempo real a partir de los datos registrados.')
add_body(doc, 'KPIs disponibles:')
kpis = [
    ('Ingredientes activos', 'total de insumos en el maestro.'),
    ('Proveedores',          'total de proveedores activos.'),
    ('Tipos de taller',      'categorías de clases configuradas.'),
    ('Sesiones registradas', 'clases ingresadas al sistema.'),
    ('Alumnos totales',      'suma de alumnos en todas las sesiones.'),
    ('Costo/alumno prom.',   'promedio ponderado de todas las sesiones.'),
    ('Costo MP total',       'suma de costos de todas las sesiones.'),
    ('Carreras',             'programas académicos configurados.'),
    ('Módulos',              'módulos académicos totales.'),
    ('Alertas stock',        'ingredientes bajo el stock mínimo definido.'),
]
for kpi, desc in kpis:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    r1 = p.add_run(f'{kpi}: ')
    set_font(r1, bold=True, size=10)
    r2 = p.add_run(desc)
    set_font(r2, size=10)

add_heading(doc, '6.2 Simulador de Costos', 2)
add_body(doc, 'Permite simular variaciones de precio por categoría sin modificar la base de datos. Ideal para anticipar el impacto de alzas de precios.')
add_steps(doc, [
    ('Acceder',   'Ve a Análisis → Simulador de Costos.'),
    ('Ajustar',   'Mueve el slider de cada categoría para simular una variación (ej: +15% en Lácteos).'),
    ('Resultado', 'El panel muestra el impacto inmediato en el costo total de sesiones.'),
    ('Resetear',  'Haz clic en Resetear para volver a los valores reales.'),
])
add_tip(doc, 'Útil antes de negociar con proveedores o revisar aranceles de matrícula.')

add_heading(doc, '6.3 Comparador de Precios', 2)
add_body(doc, 'Muestra los ingredientes agrupados por proveedor ordenados por precio. Facilita identificar quién ofrece mejor precio para cada insumo.')
add_steps(doc, [
    ('Acceder', 'Navega a Análisis → Comparador de Precios.'),
    ('Filtrar', 'Usa el selector de categoría para comparar ingredientes del mismo grupo.'),
    ('Leer',    'Los ingredientes con el precio más bajo por unidad aparecen destacados.'),
])

add_heading(doc, '6.4 Historial de Precios', 2)
add_body(doc, 'Trazabilidad de cambios de precio por ingrediente. Permite ver la variación porcentual histórica.')
add_steps(doc, [
    ('Seleccionar',  'Elige un ingrediente en el filtro superior.'),
    ('Registrar',    'Haz clic en Registrar variación para ingresar un cambio de precio.'),
    ('Datos',        'Ingresa: precio nuevo, fecha del cambio y motivo.'),
    ('Ver historia', 'La tabla muestra cada precio histórico con la variación % respecto al anterior. Alzas en rojo, bajas en verde.'),
])

add_heading(doc, '6.5 Presupuesto vs Real', 2)
add_body(doc, 'Compara el gasto planificado contra el ejecutado por tipo de taller y período.')
add_steps(doc, [
    ('Crear presupuesto','Haz clic en + Nuevo presupuesto.'),
    ('Datos',            'Selecciona taller, período (ej: MARZO 2025), monto presupuestado y número de alumnos esperados.'),
    ('Comparación',      'El sistema compara automáticamente con el gasto real de sesiones del mismo taller y período.'),
    ('Resultado',        'Ítems sobre presupuesto aparecen en rojo (overspend); bajo presupuesto en verde (underspend).'),
])

add_heading(doc, '6.6 Resumen por Docente', 2)
add_body(doc, 'Agrupa el gasto de materia prima por el docente registrado en cada sesión.')
add_steps(doc, [
    ('Acceder',    'Ve a Análisis → Resumen Docentes.'),
    ('Leer tabla', 'Columnas: docente, sesiones dictadas, alumnos totales, costo MP total, costo promedio por sesión.'),
    ('Gráfico',    'El gráfico de barras compara visualmente el gasto entre docentes.'),
])
add_tip(doc, 'El nombre del docente se toma del campo "Profesor" de cada sesión. Usa nombres consistentes para agrupar correctamente.', warn=True)

add_heading(doc, '6.7 Comparativa de Períodos', 2)
add_body(doc, 'Compara costos entre dos años académicos para identificar tendencias.')
add_steps(doc, [
    ('Seleccionar', 'Elige el Período A y el Período B en los selectores.'),
    ('Comparar',    'Haz clic en Comparar.'),
    ('Resultado',   'Se muestran: costo total de cada período, variación % (↑ rojo si sube, ↓ verde si baja) y gráfico mes a mes.'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. DOCUMENTOS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. DOCUMENTOS', 1)

add_heading(doc, '7.1 Orden de Compra', 2)
add_body(doc, 'Genera órdenes de compra formales agrupadas por proveedor, listas para imprimir o enviar.')
add_steps(doc, [
    ('Acceder',    'Ve a Documentos → Orden de Compra.'),
    ('Contenido',  'El sistema agrupa automáticamente los insumos a comprar (de la Lista de Compras) por proveedor.'),
    ('Leer',       'Cada bloque muestra: nombre del proveedor, condición de pago, listado de ítems con cantidad y precio unitario.'),
    ('Imprimir',   'Haz clic en 🖨 Imprimir PDF para generar el documento desde el diálogo de impresión del navegador.'),
])
add_tip(doc, 'Los ingredientes sin proveedor asignado no aparecerán en la orden. Asegúrate de que todos los ingredientes tengan proveedor.', warn=True)

add_heading(doc, '7.2 Informe PDF por Carrera', 2)
add_body(doc, 'Informe formal con el detalle de costos de una carrera específica, diseñado para presentaciones a dirección.')
add_steps(doc, [
    ('Acceder',    'Ve a Backup → Informe PDF por Carrera.'),
    ('Seleccionar','Elige la carrera en el selector desplegable.'),
    ('Generar',    'Haz clic en Ver informe — se abre un modal con el resumen financiero y desglose de módulos.'),
    ('PDF',        'Haz clic en 🖨 Imprimir PDF para guardar como PDF desde el diálogo del navegador.'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. HERRAMIENTAS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. HERRAMIENTAS', 1)

add_heading(doc, '8.1 Búsqueda Global', 2)
add_body(doc, 'Localiza cualquier elemento del sistema desde cualquier sección sin navegar manualmente.')
add_steps(doc, [
    ('Activar',   'Presiona Ctrl + K o haz clic en el ícono de lupa del menú lateral.'),
    ('Buscar',    'Escribe el nombre de un ingrediente, proveedor, código de clase o carrera.'),
    ('Navegar',   'Usa las teclas ↑ ↓ para moverse entre resultados.'),
    ('Seleccionar','Presiona Enter para ir directamente a la sección correspondiente.'),
    ('Cerrar',    'Presiona Esc o haz clic fuera del panel de búsqueda.'),
])

add_heading(doc, '8.2 Importar desde Excel', 2)
add_body(doc, 'Carga masiva de datos usando la plantilla oficial. Ideal para cargar el sistema por primera vez o agregar muchos registros a la vez.')
add_steps(doc, [
    ('Acceder',    'Ve a Herramientas → Importar Excel.'),
    ('Plantilla',  'Haz clic en Descargar Plantilla. El archivo descargado tiene 4 hojas: Ingredientes, Sesiones, Detalles de sesión e Inventario.'),
    ('Completar',  'Llena la plantilla siguiendo las instrucciones en cada hoja. No modifiques los encabezados.'),
    ('Cargar',     'Sube el archivo completado usando el selector de archivo.'),
    ('Vista previa','El sistema muestra una vista previa con los errores resaltados antes de importar.'),
    ('Importar',   'Haz clic en Importar todo a Supabase. Al finalizar aparece el conteo de registros importados y errores.'),
])
add_tip(doc, 'La importación no elimina datos existentes — solo agrega nuevos. Evita duplicados verificando los datos antes de importar.', warn=True)

add_heading(doc, '8.3 Backup / Exportar datos', 2)
add_body(doc, 'Tres tipos de exportación Excel más el informe PDF por carrera. Se recomienda hacer backup antes de cambios estructurales.')
add_steps(doc, [
    ('Backup completo',          'Descarga un Excel con 7 hojas: Ingredientes, Proveedores, Sesiones, Inventario, Módulos, Semanas, Carreras.'),
    ('Solo sesiones',            'Exporta la vista de costos de sesiones filtrada con cálculos de costo/alumno.'),
    ('Ingredientes y precios',   'Lista completa de ingredientes con precios, stock mínimo y proveedor asignado.'),
    ('Informe carrera',          'Ver sección 7.2.'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. CONFIGURACIÓN Y ACCESIBILIDAD
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. CONFIGURACIÓN Y ACCESIBILIDAD', 1)

add_heading(doc, 'Modo oscuro', 2)
add_body(doc, 'Haz clic en el ícono 🌙 en el pie del menú lateral para alternar entre modo claro y oscuro. La preferencia se guarda automáticamente en el navegador.')

add_heading(doc, 'Auto-actualización', 2)
add_body(doc, 'El botón Auto en el menú activa la actualización automática de todos los datos cada 60 segundos. Útil en pantallas de monitoreo permanente o salas de reunión.')

add_heading(doc, 'Instalación como aplicación (PWA)', 2)
add_body(doc,
    'El sistema es instalable como aplicación en computadores y celulares sin necesidad de la App Store:')
add_steps(doc, [
    ('Chrome / Edge', 'Aparece el ícono ⊕ en la barra de dirección. Haz clic para instalar.'),
    ('Safari (iOS)',  'Toca el botón Compartir y selecciona "Agregar a pantalla de inicio".'),
    ('Android',       'El navegador muestra una notificación de instalación automáticamente.'),
])
add_tip(doc, 'La aplicación instalada funciona como un acceso directo optimizado y puede usarse en pantalla completa.')

add_heading(doc, 'Acceso desde múltiples dispositivos', 2)
add_body(doc, 'El sistema funciona en cualquier dispositivo con navegador actualizado. Los datos son compartidos en tiempo real para todos los usuarios conectados, ya que todo se almacena en la nube (Supabase).')

# ════════════════════════════════════════════════════════════════════════════
# PIE DE PÁGINA
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sistema de Costos Gastronómicos · Instituto Culinary · Documento interno')
set_font(r, size=9, color=GRIS)

OUT = r'C:\Users\dmoov\Documents\Programas\costos-gastronomia\Manual_Procedimiento_Culinary.docx'
doc.save(OUT)
print(f'OK: {OUT}')
