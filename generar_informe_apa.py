import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r'C:\Users\dmoov\Documents\Programas\costos-gastronomia\Culinary_Informe_APA_Sistema_Costos.docx'

doc = Document()

# ── Márgenes APA: 2.54 cm (1 pulgada) en todos los lados ──────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)

# ── Fuente base: Times New Roman 12pt ─────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)   # doble espacio
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)

def set_heading_style(para, level):
    if level == 1:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
    elif level == 2:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.bold = True
    elif level == 3:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = Inches(0.5)
        for run in para.runs:
            run.bold = True
            run.italic = True

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text + '.')
    run.bold = True
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def page_break():
    doc.add_page_break()

def blank_line():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

def tabla(headers, rows, caption=None, note=None):
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(12)
        cp.paragraph_format.space_after = Pt(0)
        r1 = cp.add_run('Tabla. ')
        r1.italic = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = cp.add_run(caption)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)

    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        # Top border del encabezado
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # sin bordes laterales en APA

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    if note:
        np = doc.add_paragraph()
        np.paragraph_format.space_before = Pt(6)
        np.paragraph_format.space_after = Pt(12)
        nr1 = np.add_run('Nota. ')
        nr1.italic = True
        nr1.font.name = 'Times New Roman'
        nr1.font.size = Pt(11)
        nr2 = np.add_run(note)
        nr2.font.name = 'Times New Roman'
        nr2.font.size = Pt(11)
    return t

# ═══════════════════════════════════════════════════════
# PORTADA (APA 7ma edición)
# ═══════════════════════════════════════════════════════
for _ in range(4):
    blank_line()

p_title = doc.add_paragraph()
p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.line_spacing = Pt(24)
r = p_title.add_run('Implementación de un Sistema de Gestión de Costos de Materia Prima\npara Programas de Formación Culinaria:\nCaso Instituto Culinary, Sede Santiago')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

blank_line()
blank_line()

for txt in ['Área de Gestión Académica y Administración',
            'Instituto Culinary',
            'Departamento de Planificación y Costos',
            'Santiago, Chile']:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    r = p.add_run(txt)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

blank_line()
blank_line()

p_date = doc.add_paragraph()
p_date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.line_spacing = Pt(24)
r = p_date.add_run('Junio de 2026')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════════════════════
# RESUMEN / ABSTRACT
# ═══════════════════════════════════════════════════════
h1('Resumen')
body(
    'El presente informe describe el diseño, desarrollo e implementación del Sistema de Costos '
    'Gastronómicos del Instituto Culinary, sede Santiago. El sistema fue desarrollado como una '
    'aplicación web de página única (SPA) con almacenamiento en la nube mediante Supabase '
    '(PostgreSQL), con el objetivo de centralizar el registro, cálculo y análisis de los costos '
    'de materia prima generados en los talleres culinarios de la institución. El período de análisis '
    'corresponde al año académico 2026, compuesto por 44 semanas lectivas. Los resultados muestran '
    'un gasto total de materia prima de $7.075.872 CLP para 179 órdenes de producto (OPs) únicas, '
    'distribuidas en 1.135 registros de sesión y nueve tipos de taller activos, con 1.067 alumnos '
    'matriculados y un costo promedio por alumno de $328 CLP por sesión. '
    'El taller de mayor costo unitario fue el Taller de Técnicas y Procesos Innovadores en '
    'Pastelería (TTPIP) con $1.106 por alumno, mientras que el Taller de Panadería de Especialidad '
    '(TPE) registró el menor costo con $254 por alumno. El sistema permite filtrar, exportar e '
    'integrar datos de 613 ingredientes, 22 proveedores, 14 tipos de taller y múltiples estructuras '
    'curriculares. Se identificaron cinco talleres sin registros de recetas, los cuales representan '
    'una oportunidad de mejora para completar la trazabilidad del sistema.'
)
blank_line()
p_kw = doc.add_paragraph()
p_kw.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_kw.paragraph_format.first_line_indent = Inches(0.5)
p_kw.paragraph_format.line_spacing = Pt(24)
kw_label = p_kw.add_run('Palabras clave: ')
kw_label.italic = True
kw_label.font.name = 'Times New Roman'
kw_label.font.size = Pt(12)
kw_rest = p_kw.add_run('costos gastronómicos, gestión de costos educativos, materia prima, sistema de información, formación culinaria')
kw_rest.font.name = 'Times New Roman'
kw_rest.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════════════════════
# CUERPO
# ═══════════════════════════════════════════════════════
h1('Introducción')
body(
    'La gestión eficiente de los costos de insumos en instituciones de educación gastronómica representa '
    'un desafío constante para los equipos de administración académica. A diferencia de los entornos '
    'de restauración comercial, donde el costeo es una práctica estandarizada y vinculada directamente '
    'a la rentabilidad, en el ámbito formativo los costos de materia prima se distribuyen entre '
    'múltiples secciones, fechas y grupos de alumnos, lo que dificulta su seguimiento y análisis.'
)
body(
    'El Instituto Culinary, institución de formación culinaria con sede en Santiago de Chile, '
    'identificó la necesidad de contar con una herramienta digital que permitiera centralizar '
    'el registro de ingredientes utilizados en sus talleres prácticos, calcular automáticamente '
    'los costos asociados y generar reportes para la toma de decisiones académico-administrativas.'
)
body(
    'En respuesta a esta necesidad, se desarrolló el Sistema de Costos Gastronómicos, una aplicación '
    'web que integra una base de datos relacional en la nube con una interfaz de usuario moderna y '
    'accesible desde cualquier dispositivo. El presente informe documenta los objetivos, la '
    'arquitectura técnica, los datos procesados y los principales hallazgos derivados del uso del '
    'sistema durante el período académico 2026.'
)

h1('Objetivos')
h2('Objetivo General')
body(
    'Desarrollar e implementar un sistema digital de gestión de costos de materia prima para los '
    'talleres culinarios del Instituto Culinary, sede Santiago, que permita el registro sistemático, '
    'el cálculo automatizado y el análisis visual de los gastos de insumos durante el período '
    'académico 2026.'
)
h2('Objetivos Específicos')
for obj in [
    'Centralizar el maestro de ingredientes con precios actualizados, unidades de medida estandarizadas y categorización por tipo de insumo.',
    'Registrar las recetas utilizadas en cada sesión de taller, asociando cantidades e ingredientes a las sesiones correspondientes del calendario académico.',
    'Calcular automáticamente el costo total de materia prima por sesión y el costo promedio por alumno.',
    'Generar reportes visuales e indicadores clave de desempeño (KPIs) para la gestión académica y presupuestaria.',
    'Facilitar la exportación de datos para auditorías, informes institucionales y carga masiva de información.',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.line_spacing = Pt(24)
    r = p.add_run(obj)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

h1('Marco Teórico')
h2('Costeo de Materia Prima en Educación Culinaria')
body(
    'El costo de materia prima (CMP) en contextos de formación culinaria se define como el valor '
    'monetario de los ingredientes consumidos en las actividades prácticas de aprendizaje. A '
    'diferencia del costo en restaurantes, donde el CMP se relaciona directamente con el precio '
    'de venta al público, en la educación el CMP es un costo de operación que debe ser planificado '
    'y controlado dentro del presupuesto institucional (Labensky et al., 2019).'
)
body(
    'La literatura en administración de servicios de alimentación señala que una gestión adecuada '
    'del CMP requiere tres elementos fundamentales: (a) un maestro de ingredientes con precios '
    'actualizados y unidades de medida estandarizadas; (b) recetas estandarizadas que especifiquen '
    'con precisión las cantidades requeridas por preparación; y (c) un sistema de seguimiento que '
    'permita comparar el costo planificado con el costo real (Ninemeier & Hayes, 2006).'
)
h2('Sistemas de Información para Gestión de Costos')
body(
    'Los sistemas de información de gestión (MIS, por sus siglas en inglés) han evolucionado '
    'significativamente con la adopción de tecnologías en la nube. En el contexto de pequeñas y '
    'medianas instituciones educativas, los sistemas web modernos basados en arquitecturas de '
    'página única (SPA) y bases de datos como servicio (DBaaS) representan una alternativa '
    'eficiente y de bajo costo frente a soluciones ERP corporativas (O\'Brien & Marakas, 2020).'
)
body(
    'Supabase, la plataforma utilizada como backend del sistema, es una solución de código abierto '
    'que proporciona una API REST automática sobre PostgreSQL, permitiendo operaciones CRUD sin '
    'necesidad de desarrollar un servidor backend independiente. Esta arquitectura reduce '
    'significativamente los tiempos y costos de desarrollo (Supabase, 2024).'
)

h1('Metodología')
h2('Diseño del Sistema')
body(
    'El sistema fue desarrollado siguiendo un modelo de desarrollo iterativo. La primera fase '
    'consistió en el levantamiento de requerimientos con el equipo de coordinación académica '
    'del Instituto Culinary, identificando las entidades principales: ingredientes, proveedores, '
    'sesiones de clase, tipos de taller, módulos académicos y semanas del calendario lectivo.'
)
body(
    'La arquitectura elegida corresponde a un modelo de tres capas: (a) capa de presentación: '
    'archivo HTML único con JavaScript vanilla y Chart.js para visualizaciones; (b) capa de '
    'datos: base de datos PostgreSQL gestionada por Supabase con PostgREST como capa de API; '
    'y (c) capa de despliegue: PythonAnywhere para el frontend y los servicios en la nube de '
    'Supabase para el backend.'
)
h2('Procesamiento de Datos')
body(
    'Los datos iniciales fueron extraídos del archivo "Dashboard Culinary Depurado Santiago.xlsx", '
    'que contenía cuatro hojas: ingredientes con precios de proveedor, sesiones del calendario '
    'académico, detalle de ingredientes por sesión, y estructura de talleres y secciones. '
    'Se desarrollaron scripts en Python para la carga automatizada vía API REST de Supabase.'
)
body(
    'Se aplicó un proceso de conversión de unidades para estandarizar los precios: los insumos '
    'comercializados en gramos (GR) o centímetros cúbicos (CC) fueron convertidos a kilogramos '
    '(kg) o litros (L) multiplicando el precio unitario por 1.000, cuando el formato del envase '
    'era menor a 1.000 unidades. Adicionalmente, se aplicaron factores de conversión peso-unidad '
    'para 77 ingredientes que se venden por peso pero se utilizan por pieza (e.g., COLAPEZ: '
    '0,002 kg/hoja; huevos: precio por unidad).'
)
h2('Modelo de Costeo')
body(
    'El modelo de costeo adoptado calcula el costo de materia prima por sesión como la suma del '
    'producto entre la cantidad utilizada de cada ingrediente y su precio unitario estandarizado. '
    'La fórmula es la siguiente:'
)
p_formula = doc.add_paragraph()
p_formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_formula.paragraph_format.line_spacing = Pt(24)
p_formula.paragraph_format.space_before = Pt(12)
p_formula.paragraph_format.space_after = Pt(12)
r = p_formula.add_run('CMP_sesión = Σ (cantidad_utilizada_i × precio_unitario_i)')
r.font.name = 'Courier New'
r.font.size = Pt(11)
r.bold = True

body(
    'La cantidad utilizada corresponde a la cantidad de la receta tal como está registrada, sin '
    'multiplicar por el número de alumnos de la sección. Este criterio fue validado con el Excel '
    'de referencia, verificándose una concordancia del 100% entre los totales calculados por el '
    'sistema y los registrados en la planilla de origen.'
)
body(
    'El costo por alumno se calcula dividiendo el costo total de la sesión por el número de '
    'alumnos matriculados en esa sección específica, lo que permite comparar la eficiencia '
    'de costo entre talleres con distintos tamaños de grupo.'
)

h1('Resultados')
h2('Estadísticas Generales del Período 2026')
body(
    'Durante el período académico 2026 (marzo a diciembre), el sistema registró un total de '
    '179 órdenes de producto (OPs) únicas, las cuales generan 1.135 registros de sesión distribuidos '
    'en 44 semanas lectivas. El total de alumnos matriculados asciende a 1.067, calculado como la '
    'suma de alumnos por combinación única de taller y sección. El gasto total de materia prima '
    'alcanzó los $7.075.872 CLP, con un costo promedio por alumno por sesión de $328 CLP.'
)

tabla(
    ['Indicador', 'Valor'],
    [
        ['Total OPs únicas', '179'],
        ['Total sesiones registradas', '1.135'],
        ['Total alumnos matriculados', '1.067'],
        ['Total semanas lectivas', '44'],
        ['Total ingredientes en maestro', '613'],
        ['Total proveedores activos', '22'],
        ['Tipos de taller', '14 (9 con recetas)'],
        ['Registros ingrediente-sesión', '23.213'],
        ['Gasto total materia prima', '$7.075.872 CLP'],
        ['Costo promedio por alumno/sesión', '$328 CLP'],
    ],
    caption='Resumen de indicadores generales del Sistema de Costos Gastronómicos, Sede Santiago 2026',
    note='Los valores corresponden al período académico marzo–diciembre 2026. El costo promedio por alumno/sesión se calcula como el promedio de los cocientes individuales de cada sesión, no como la división del total.'
)

h2('Análisis por Categoría de Ingrediente')
body(
    'El análisis de costos por categoría de ingrediente revela que los abarrotes (harinas, aceites, '
    'azúcar, sal y productos secos en general) constituyen el mayor gasto, representando el 28,8% '
    'del total. Le siguen los lácteos (17,5%) y el chocolate (11,6%), lo que refleja la orientación '
    'del programa hacia la pastelería y repostería. La Tabla 2 presenta el detalle por categoría.'
)

tabla(
    ['Categoría', 'Gasto Total (CLP)', 'Participación (%)'],
    [
        ['Abarrote',         '$2.040.961',  '28,8%'],
        ['Lácteo',           '$1.241.601',  '17,5%'],
        ['Chocolate',        '$824.444',    '11,6%'],
        ['Fruta y Verdura',  '$705.926',    '10,0%'],
        ['Pescado/Marisco',  '$389.899',     '5,5%'],
        ['Vacuno',           '$346.539',     '4,9%'],
        ['Huevo',            '$332.590',     '4,7%'],
        ['Congelado',        '$303.393',     '4,3%'],
        ['Art. de Cocina',   '$273.737',     '3,9%'],
        ['Resto (25 cat.)',  '$617.278',     '8,7%'],
        ['TOTAL',            '$7.075.872',  '100,0%'],
    ],
    caption='Gasto de materia prima por categoría de ingrediente, período 2026',
    note='La categoría "Resto" agrupa 25 categorías adicionales con participaciones individuales inferiores al 3,0%, incluyendo cerdo, pollo, panadería, embutidos, vinos, licores y otros.'
)

h2('Análisis por Tipo de Taller')
body(
    'El análisis por tipo de taller permite identificar diferencias significativas en el costo '
    'de materia prima entre las distintas especialidades. El Taller Introductorio de Pastelería '
    '(TIP) registra el mayor costo total absoluto ($1.119.984), mientras que el Taller de '
    'Técnicas y Procesos Innovadores en Pastelería (TTPIP) presenta el mayor costo por alumno '
    '($1.106), lo que refleja el uso de ingredientes de mayor valor en esta especialización.'
)

tabla(
    ['Código', 'Taller', 'Sesiones', 'Costo Total (CLP)', 'CPA (CLP)'],
    [
        ['TIP',   'Introductorio Pastelería',          '140', '$1.119.984', '$471'],
        ['TBD',   'Banquetería Dulce',                  '64', '$1.062.113', '$909'],
        ['TCCHO', 'Confitería y Chocolatería',           '64',   '$975.967', '$744'],
        ['TCCM',  'Cocina Chilena y del Mundo',         '100',   '$901.539', '$440'],
        ['TTPIP', 'Técnicas Innovadoras Pastelería',     '40',   '$752.045', '$1.106'],
        ['TIC',   'Introductorio Cocina',               '155',   '$723.792', '$275'],
        ['TBS',   'Banquetería Salada',                  '96',   '$657.917', '$376'],
        ['TTPIC', 'Técnicas Innovadoras Cocina',         '40',   '$486.289', '$715'],
        ['TPE',   'Panadería Especialidad',              '76',   '$396.225', '$254'],
        ['TOTAL', '9 talleres activos',                 '775', '$7.075.872',   '$328*'],
    ],
    caption='Costo de materia prima por tipo de taller activo, período 2026',
    note='* El CPA total es el promedio ponderado. CPA = Costo por Alumno por sesión. Cinco talleres (TMA, TAS, TIR, TMI, TEAC) tienen 360 sesiones registradas sin recetas cargadas, por lo que su costo no está incluido en este análisis.'
)

h2('Distribución Mensual del Gasto')
body(
    'La distribución temporal del gasto muestra variaciones mensuales que responden al calendario '
    'académico y a la concentración de determinadas asignaturas en ciertos períodos. Mayo registra '
    'el mayor gasto ($1.178.684), seguido de marzo ($1.014.475) y junio ($1.076.417). Julio '
    'presenta el menor gasto ($348.671), posiblemente asociado a un período de menor actividad '
    'académica o pausa invernal.'
)

tabla(
    ['Mes', 'Gasto (CLP)', 'Participación (%)'],
    [
        ['Marzo',      '$1.014.475', '14,3%'],
        ['Abril',        '$965.449', '13,6%'],
        ['Mayo',       '$1.178.684', '16,7%'],
        ['Junio',      '$1.076.417', '15,2%'],
        ['Julio',        '$348.671',  '4,9%'],
        ['Agosto',       '$528.122',  '7,5%'],
        ['Septiembre',   '$593.151',  '8,4%'],
        ['Octubre',      '$458.301',  '6,5%'],
        ['Noviembre',    '$588.237',  '8,3%'],
        ['Diciembre',    '$324.365',  '4,6%'],
        ['TOTAL',      '$7.075.872', '100,0%'],
    ],
    caption='Distribución mensual del gasto de materia prima, período 2026',
    note='Los valores reflejan el gasto real registrado en el sistema según las recetas ingresadas por sesión.'
)

h1('Discusión')
body(
    'Los resultados obtenidos permiten establecer un punto de referencia cuantitativo para la '
    'gestión presupuestaria de los talleres culinarios del Instituto Culinary. El costo promedio '
    'de $328 CLP por alumno por sesión constituye una métrica útil para la planificación de '
    'presupuestos futuros y para la evaluación de la eficiencia en el uso de insumos.'
)
body(
    'La concentración del gasto en abarrotes, lácteos y chocolates (57,9% del total) sugiere que '
    'las estrategias de reducción de costos deberían focalizarse en estos grupos. Negociaciones '
    'con proveedores para compras por volumen, el uso de productos sustitutos o la optimización '
    'de recetas podrían tener un impacto significativo en el costo total de operación.'
)
body(
    'La disparidad entre el costo por alumno de los talleres avanzados (TTPIP: $1.106, TBD: $909) '
    'y los talleres introductorios (TIC: $275, TPE: $254) es esperable dado el nivel de '
    'especialización de los ingredientes. Sin embargo, este análisis permite identificar los '
    'talleres donde el costo de insumos podría representar una barrera para la sustentabilidad '
    'del programa, especialmente si los aranceles no reflejan estos costos diferenciados.'
)
body(
    'Una limitación relevante del análisis actual es que cinco tipos de taller (TMA, TAS, TIR, '
    'TMI y TEAC), que representan el 31,7% de las sesiones totales (360 de 1.135), no tienen '
    'recetas registradas en el sistema. Esto impide una visión completa del gasto institucional '
    'y subestima el costo total real. La incorporación de estos registros es la prioridad más '
    'urgente para mejorar la representatividad del sistema.'
)

h1('Conclusiones')
body(
    'El Sistema de Costos Gastronómicos del Instituto Culinary representa un avance significativo '
    'en la gestión de información académico-administrativa de la institución. Su implementación '
    'permitió centralizar datos que antes se manejaban de forma dispersa en planillas individuales, '
    'generando un registro único, auditable y actualizable en tiempo real.'
)
body(
    'Los principales logros del sistema son: (a) la carga completa del período 2026 con 179 OPs '
    'únicas, 1.135 registros de sesión y 23.213 registros de ingredientes-sesión; (b) el cálculo '
    'automático y consistente del costo por sesión y por alumno, validado contra el Excel de '
    'referencia con una concordancia del 100%; y (c) la disponibilidad de un dashboard interactivo '
    'con filtros, gráficos y exportación de datos para la toma de decisiones.'
)
body(
    'Como trabajo futuro, se recomienda: completar la carga de recetas para los cinco talleres '
    'sin datos; vincular las sesiones con las semanas del calendario académico para habilitar '
    'análisis por semana; incorporar un módulo de presupuesto vs. real; y considerar la extensión '
    'del sistema a otras sedes del Instituto para generar comparativas institucionales.'
)

h1('Referencias')
referencias = [
    'Labensky, S. R., Hause, A. M., & Labensky, S. R. (2019). *On cooking: A textbook of culinary fundamentals* (6.ª ed.). Pearson.',
    'Ninemeier, J. D., & Hayes, D. K. (2006). *Restaurant operations management: Principles and practices*. Pearson Prentice Hall.',
    "O'Brien, J. A., & Marakas, G. M. (2020). *Introduction to information systems* (16.ª ed.). McGraw-Hill.",
    'Supabase. (2024). *Supabase documentation: PostgREST API*. https://supabase.com/docs/guides/api',
    'Instituto Culinary. (2026). *Dashboard Culinary Depurado Santiago.xlsx* [Archivo de datos internos]. Santiago, Chile: Instituto Culinary.',
]
for ref in referencias:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Procesar markdown italics (* ... *)
    parts = ref.split('*')
    for j, part in enumerate(parts):
        r = p.add_run(part)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        if j % 2 == 1:
            r.italic = True

# ── Encabezado APA (running head) ─────────────────────────────
from docx.oxml.ns import qn as ns_qn
section2 = doc.sections[0]
header = section2.header
header.paragraphs[0].clear()
p_h = header.paragraphs[0]
p_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rh = p_h.add_run('SISTEMA DE COSTOS GASTRONÓMICOS — INSTITUTO CULINARY')
rh.font.name = 'Times New Roman'
rh.font.size = Pt(12)

# Pie de página con número
footer = section2.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fpr = fp.add_run()
fpr.font.name = 'Times New Roman'
fpr.font.size = Pt(12)
fld = OxmlElement('w:fldChar')
fld.set(qn('w:fldCharType'), 'begin')
fpr._r.append(fld)
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fpr._r.append(instrText)
fld2 = OxmlElement('w:fldChar')
fld2.set(qn('w:fldCharType'), 'end')
fpr._r.append(fld2)

doc.save(OUT)
print(f'Informe APA guardado: {OUT}')
