import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\dmoov\Documents\Programas\costos-gastronomia\Culinary_Guia_Presentacion.docx'

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)

base = doc.styles['Normal']
base.font.name = 'Calibri'
base.font.size = Pt(11)

def hexcolor(h):
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

GOLD   = hexcolor('fbbf24')
DARK   = hexcolor('1e1e1e')
MUTED  = hexcolor('6b7280')
GREEN  = hexcolor('22c55e')
RED    = hexcolor('ef4444')
BLUE   = hexcolor('3b82f6')
WHITE  = hexcolor('ffffff')

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(16)
    r.font.color.rgb = color or GOLD
    return p

def h2(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.color.rgb = color or DARK
    return p

def body(text, indent=False, muted=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    if muted:
        r.font.color.rgb = MUTED
    return p

def bullet(text, bold_prefix=None, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' ')
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        if color:
            r1.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('─' * 90)
    r.font.name = 'Calibri'
    r.font.size = Pt(8)
    r.font.color.rgb = hexcolor('d1d5db')
    return p

def slide_block(num, titulo, duracion, objetivo, puntos_clave, consejo=None, datos=None):
    """Crea un bloque de slide con tabla informativa."""
    t = doc.add_table(rows=1, cols=12)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Número de slide
    c0 = t.rows[0].cells[0]
    c0.merge(t.rows[0].cells[1])
    set_cell_bg(c0, '1e1e1e')
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f'SLIDE\n{num:02d}')
    r0.bold = True; r0.font.name = 'Calibri'; r0.font.size = Pt(14)
    r0.font.color.rgb = GOLD
    # Título del slide
    c1 = t.rows[0].cells[2]
    c1.merge(t.rows[0].cells[8])
    set_cell_bg(c1, '141414')
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(titulo)
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(13)
    r1.font.color.rgb = WHITE
    # Duración
    c2 = t.rows[0].cells[9]
    c2.merge(t.rows[0].cells[11])
    set_cell_bg(c2, '292929')
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f'⏱ {duracion}')
    r2.font.name = 'Calibri'; r2.font.size = Pt(10)
    r2.font.color.rgb = GOLD

    # Objetivo
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.left_indent = Inches(0.2)
    p_obj.paragraph_format.space_before = Pt(4)
    p_obj.paragraph_format.space_after = Pt(2)
    r_lbl = p_obj.add_run('OBJETIVO: ')
    r_lbl.bold = True; r_lbl.font.name = 'Calibri'; r_lbl.font.size = Pt(10)
    r_lbl.font.color.rgb = BLUE
    r_txt = p_obj.add_run(objetivo)
    r_txt.font.name = 'Calibri'; r_txt.font.size = Pt(10)

    # Puntos clave
    p_kl = doc.add_paragraph()
    p_kl.paragraph_format.left_indent = Inches(0.2)
    p_kl.paragraph_format.space_before = Pt(2)
    p_kl.paragraph_format.space_after = Pt(1)
    r_kl = p_kl.add_run('PUNTOS CLAVE A DESTACAR:')
    r_kl.bold = True; r_kl.font.name = 'Calibri'; r_kl.font.size = Pt(10)
    r_kl.font.color.rgb = MUTED
    for pk in puntos_clave:
        p_pk = doc.add_paragraph()
        p_pk.paragraph_format.left_indent = Inches(0.5)
        p_pk.paragraph_format.space_before = Pt(1)
        p_pk.paragraph_format.space_after = Pt(1)
        r_arr = p_pk.add_run('→  ')
        r_arr.font.name = 'Calibri'; r_arr.font.size = Pt(10)
        r_arr.font.color.rgb = GOLD
        r_pk = p_pk.add_run(pk)
        r_pk.font.name = 'Calibri'; r_pk.font.size = Pt(10)

    if datos:
        p_d = doc.add_paragraph()
        p_d.paragraph_format.left_indent = Inches(0.2)
        p_d.paragraph_format.space_before = Pt(2)
        r_dl = p_d.add_run('CIFRAS CLAVE:  ')
        r_dl.bold = True; r_dl.font.name = 'Calibri'; r_dl.font.size = Pt(10)
        r_dl.font.color.rgb = GREEN
        r_dv = p_d.add_run(datos)
        r_dv.font.name = 'Calibri'; r_dv.font.size = Pt(10)

    if consejo:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.left_indent = Inches(0.2)
        p_c.paragraph_format.space_before = Pt(2)
        p_c.paragraph_format.space_after = Pt(6)
        r_cl = p_c.add_run('💡 CONSEJO:  ')
        r_cl.bold = True; r_cl.font.name = 'Calibri'; r_cl.font.size = Pt(10)
        r_cl.font.color.rgb = GOLD
        r_cv = p_c.add_run(consejo)
        r_cv.italic = True; r_cv.font.name = 'Calibri'; r_cv.font.size = Pt(10)

    divider()

# ═══════════════════════════════════════════════════════
# PORTADA GUÍA
# ═══════════════════════════════════════════════════════
p_port = doc.add_paragraph()
p_port.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_port.paragraph_format.space_before = Pt(30)
r = p_port.add_run('GUÍA DE PRESENTACIÓN')
r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(28)
r.font.color.rgb = GOLD

p_sub = doc.add_paragraph()
p_sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Sistema de Costos Gastronómicos · Instituto Culinary')
r2.font.name = 'Calibri'; r2.font.size = Pt(16); r2.font.color.rgb = MUTED

doc.add_paragraph()
p_inf = doc.add_paragraph()
p_inf.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_inf.add_run('Esta guía acompaña la presentación "Culinary_Sistema_Costos_Presentacion.pptx"\nSantiago, Junio 2026  ·  13 slides  ·  Duración estimada: 25–35 minutos')
r3.font.name = 'Calibri'; r3.font.size = Pt(11); r3.font.color.rgb = MUTED

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# SECCIÓN 1: PREPARACIÓN
# ═══════════════════════════════════════════════════════
h1('1.  Preparación Previa a la Presentación')

h2('Materiales necesarios')
for item in [
    'Archivo PPT: Culinary_Sistema_Costos_Presentacion.pptx',
    'Acceso a internet para demo en vivo (opcional): culinary.pythonanywhere.com',
    'Informe APA impreso o digital: Culinary_Informe_APA_Sistema_Costos.docx',
    'Proyector o pantalla con resolución mínima 1280×720 (formato 16:9)',
    'Puntero láser o dispositivo de control remoto de diapositivas (recomendado)',
]:
    bullet(item)

h2('Configuración técnica')
body('Antes de comenzar, verificar:')
for item in [
    'Abrir el PPT y poner en modo presentación (F5 o Vista → Presentación con diapositivas).',
    'Revisar que los colores oscuros del fondo se visualicen correctamente en el proyector.',
    'Si se hará demo en vivo: tener abierta la pestaña del sistema en el navegador con la sesión iniciada.',
    'Configurar el Panel de notas (Vista → Panel de notas) para usar esta guía durante la exposición.',
]:
    bullet(item)

h2('Tiempo estimado por sección')
t = doc.add_table(rows=6, cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = [('Sección', '3C9EFF'), ('Slides', '3C9EFF'), ('Duración', '3C9EFF')]
data = [
    ('Introducción y contexto', '1–3', '4–5 min'),
    ('La solución y arquitectura', '4–5', '4–5 min'),
    ('Datos y KPIs', '6–7', '5–7 min'),
    ('Análisis de costos', '8–11', '8–10 min'),
    ('Hallazgos y cierre', '12–13', '4–5 min'),
]
for ci, (h, col) in enumerate(headers):
    cell = t.rows[0].cells[ci]
    set_cell_bg(cell, '1e1e1e')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = GOLD
for ri, (sec, sl, dur) in enumerate(data):
    row = t.rows[ri+1]
    for ci, val in enumerate([sec, sl, dur]):
        cell = row.cells[ci]
        set_cell_bg(cell, '141414' if ri % 2 == 0 else '1a1a1a')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.name = 'Calibri'; r.font.size = Pt(11)
        r.font.color.rgb = WHITE if ci == 0 else GOLD
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# SECCIÓN 2: GUÍA SLIDE A SLIDE
# ═══════════════════════════════════════════════════════
h1('2.  Guía Slide por Slide')
body('Cada bloque indica el objetivo del slide, los puntos clave a mencionar, cifras de apoyo y consejos para el presentador.')
divider()

slide_block(
    1, 'PORTADA — Sistema de Costos Gastronómicos', '1–2 min',
    'Generar la primera impresión correcta: el sistema tiene identidad institucional, es profesional y pertenece al Instituto Culinary.',
    [
        'Presentar el nombre completo del sistema y el campus al que corresponde.',
        'Mencionar que es una herramienta de uso interno para gestión académico-administrativa.',
        'Indicar el período de análisis: año académico 2026.',
    ],
    consejo='No leer el slide. Presentar brevemente quién habla, en qué contexto y cuánto durará la exposición. Crear expectativa.'
)

slide_block(
    2, 'AGENDA — Contenido de la presentación', '1 min',
    'Dar al público una hoja de ruta clara para que puedan seguir el hilo de la presentación.',
    [
        'Recorrer brevemente los 9 puntos del índice.',
        'Señalar que se mostrará demo en vivo del sistema si el tiempo lo permite.',
        'Indicar cuándo habrá espacio para preguntas (al final).',
    ],
    consejo='No explicar cada punto en detalle. Solo nombrarlos y avanzar. El tiempo es valioso.'
)

slide_block(
    3, 'CONTEXTO — Problema y necesidad', '3–4 min',
    'Que el público entienda por qué se desarrolló el sistema: había un vacío real en la gestión de información.',
    [
        'Falta de visibilidad: los costos no eran comparables entre talleres ni entre períodos.',
        'Cálculos manuales en Excel: propensos a errores, sin trazabilidad ni historial.',
        'Sin datos para tomar decisiones presupuestarias informadas.',
        'La necesidad era real: la institución no podía responder "¿cuánto cuesta por alumno en TTPIP?"',
    ],
    consejo='Hacer énfasis en el impacto práctico. Si el público es directivo, destacar la falta de datos para planificación presupuestaria. Si es técnico, destacar la falta de trazabilidad.'
)

slide_block(
    4, 'LA SOLUCIÓN — Funciones del sistema', '3 min',
    'Mostrar qué hace el sistema de forma clara y concisa.',
    [
        'Es una plataforma web: no requiere instalación, funciona desde cualquier navegador.',
        'Las 6 funciones principales están integradas en una sola herramienta.',
        'El sistema centraliza todo: ingredientes, sesiones, carreras, compras e informes.',
        'Mencionar que permite exportar a Excel y PDF para reportes institucionales.',
    ],
    consejo='Puedes abrir el sistema en vivo aquí si hay conectividad. Mostrar el dashboard brevemente refuerza mucho el mensaje.'
)

slide_block(
    5, 'ARQUITECTURA — Stack tecnológico', '2–3 min',
    'Generar confianza técnica: el sistema usa tecnologías modernas, estables y de bajo costo.',
    [
        'Frontend: un solo archivo HTML. No requiere servidor de aplicaciones complejo.',
        'Backend: Supabase (PostgreSQL en la nube). API automática, sin desarrollo adicional.',
        'Despliegue: PythonAnywhere. Costo mínimo, disponible 24/7.',
        'Toda la lógica de cálculo está en vistas SQL: rápida, consistente, auditable.',
    ],
    datos='Costo estimado de infraestructura: < USD 50/mes. Sin costos de desarrollo de backend.',
    consejo='Para audiencias no técnicas, simplificar: "Es una aplicación web en la nube, segura y accesible desde cualquier dispositivo." Para audiencias técnicas, profundizar en la arquitectura SPA + PostgREST.'
)

slide_block(
    6, 'DATOS DEL PERÍODO — Carga real 2026', '3–4 min',
    'Dimensionar el volumen de información que gestiona el sistema y generar confianza en la calidad de los datos.',
    [
        '613 ingredientes con precio, unidad y categoría estandarizados.',
        '179 OPs únicas del año académico, generando 1.135 registros de sesión con recetas detalladas.',
        '23.213 registros de ingrediente por sesión: la granularidad del análisis.',
        'Todos los datos provienen del Excel operacional del Instituto, validados al 100%.',
        '44 semanas del calendario 2026 programadas en el sistema.',
    ],
    datos='23.213 registros de recetas → 7.075.872 CLP gasto total calculado automáticamente.',
    consejo='Destacar la validación: los totales del sistema coinciden exactamente con el Excel de referencia. Esto genera confianza en la calidad del dato.'
)

slide_block(
    7, 'DASHBOARD — KPIs principales', '4–5 min',
    'Comunicar los resultados más importantes del análisis: el costo total y el costo por alumno.',
    [
        'Gasto total materia prima: $7.075.872 para el año 2026.',
        'Costo promedio por alumno por sesión: $328 CLP.',
        'El gráfico de barras mensual muestra la distribución: mayo es el mes más costoso.',
        'Julio tiene el menor gasto, posiblemente por pausa académica.',
        'El dashboard permite filtrar por taller, mes, semana y año en tiempo real.',
    ],
    datos='$7.075.872 total  ·  $328 CPA promedio  ·  1.067 alumnos matriculados  ·  mayo: $1.178.684 (16,7%)',
    consejo='Este es el slide más importante de la presentación. Hablar despacio. Dar tiempo para que el público procese las cifras. Si hay preguntas aquí, es una buena señal.'
)

slide_block(
    8, 'ANÁLISIS — Costos por categoría', '3–4 min',
    'Mostrar en qué tipos de ingredientes se concentra el gasto y qué oportunidades de optimización existen.',
    [
        'Abarrote (28,8%) + Lácteo (17,5%) + Chocolate (11,6%) = 57,9% del gasto total.',
        'Los tres primeros grupos son el foco natural de negociación con proveedores.',
        'La alta participación de chocolate refleja el énfasis del programa en pastelería premium.',
        'Las 25 categorías restantes representan solo el 8,7%: impacto individual menor.',
    ],
    datos='Top 3 categorías concentran el 57,9% del gasto.',
    consejo='Las barras horizontales permiten comparar visualmente. Señalar con el puntero las barras de ABARROTE y LACTEO para reforzar el mensaje sobre optimización.'
)

slide_block(
    9, 'ANÁLISIS — Costos por tipo de taller', '4–5 min',
    'Comparar el desempeño de costo entre talleres para identificar los de mayor y menor eficiencia en uso de insumos.',
    [
        'TIP (Introductorio Pastelería) tiene el mayor costo total: $1.119.984 en 140 sesiones.',
        'TTPIP tiene el mayor costo POR ALUMNO: $1.106, casi el doble del promedio.',
        'TPE y TIC son los más eficientes por alumno ($254 y $275 respectivamente).',
        'La diferencia de CPA entre talleres avanzados e introductorios es esperable y legítima.',
        'Cinco talleres no tienen recetas → su costo real es desconocido todavía.',
    ],
    datos='Rango de CPA: $254 (TPE) → $1.106 (TTPIP). Promedio general: $328.',
    consejo='Usar la escala de colores verde/amarillo/rojo del slide para guiar la interpretación. Remarcar que el costo alto de TTPIP no es necesariamente malo, sino propio de una especialización premium.'
)

slide_block(
    10, 'TOP 5 SESIONES — Más costosas del período', '2–3 min',
    'Mostrar ejemplos concretos de las sesiones con mayor gasto, para dar contexto al análisis.',
    [
        'Todas las top 5 pertenecen a TTPIC y TBS: talleres de técnicas avanzadas.',
        'El costo total por sesión de TTPICSES3 es de $27.846, con 16–19 alumnos.',
        'El CPA de estas sesiones ($1.295–$1.740) está muy por sobre el promedio.',
        'Estos datos permiten prever el gasto de sesiones similares en futuros períodos.',
    ],
    datos='Sesión más cara: TTPICSES3 → $27.846 (CPA $1.740 con 16 alumnos).',
    consejo='Este slide es especialmente útil para la dirección académica. La granularidad del dato ($27.846 por sesión específica) demuestra el valor real del sistema.'
)

slide_block(
    11, 'TOP 10 INGREDIENTES — Por gasto acumulado', '2–3 min',
    'Identificar los ingredientes que más impactan en el costo total y que son el foco de gestión de compras.',
    [
        'Mantequilla sin sal es el ingrediente más usado en costo: $341.058.',
        'Los lácteos dominan el top: mantequilla, crema y leche en posiciones 1, 3 y 5.',
        'El huevo es el ingrediente de mayor volumen en unidades.',
        'El couverture de chocolate aparece dos veces (chocolate y couverture oscuro).',
        'Estos 10 ingredientes explican una parte desproporcionada del gasto total.',
    ],
    datos='Top 10 ingredientes: $2.172.094 estimados (30,7% del gasto total).',
    consejo='Este slide conecta directamente con la gestión de compras. Un acuerdo con el proveedor de lácteos y mantequilla puede tener más impacto que optimizar decenas de ingredientes de bajo uso.'
)

slide_block(
    12, 'HALLAZGOS — Observaciones y conclusiones', '3–4 min',
    'Sintetizar los aprendizajes más importantes del análisis para que el público los lleve consigo.',
    [
        'Hallazgo 1 (verde): el CPA promedio de $328 es razonable y coherente con el mercado.',
        'Hallazgo 2 (amarillo): TTPIP y TBD son los talleres de mayor costo unitario → monitoreo recomendado.',
        'Hallazgo 3 (azul): mayo concentra el 16,7% del gasto anual → planificación de compras.',
        'Hallazgo 4 (gris): 5 talleres sin recetas = costo real subestimado → prioridad de acción.',
    ],
    consejo='No necesitas leer los hallazgos. El color y el ícono ya lo dicen. Puedes pausar en cada uno y preguntar al público si generan alguna pregunta o reflexión.'
)

slide_block(
    13, 'PRÓXIMOS PASOS — Roadmap', '2–3 min',
    'Dejar al público con una visión clara de hacia dónde va el sistema y qué acciones son prioritarias.',
    [
        'Corto plazo: completar recetas de los 5 talleres sin datos (TMA, TAS, TIR, TMI, TEAC).',
        'Mediano plazo: módulo de presupuesto vs. real y exportación automática de informes.',
        'Largo plazo: extensión a otras sedes para comparativas institucionales.',
        'El sistema ya está operativo: los próximos pasos son completar datos, no construir.',
    ],
    consejo='Cerrar con energía y confianza. Recordar la URL del sistema: culinary.pythonanywhere.com. Invitar a explorar el sistema y proporcionar credenciales de acceso si corresponde.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# SECCIÓN 3: PREGUNTAS FRECUENTES
# ═══════════════════════════════════════════════════════
h1('3.  Preguntas Frecuentes y Cómo Responderlas')
body('Estas son las preguntas más probables según el perfil del público.')

faqs = [
    ('¿Por qué el gasto total es tan bajo? ¿No falta data?',
     'El sistema carga las recetas tal como están en el Excel operacional. Cinco talleres aún no tienen recetas ingresadas (TMA, TAS, TIR, TMI, TEAC), lo que representa 360 sesiones sin costo calculado. Si incluyéramos estimaciones de esos talleres, el total podría ser entre 35–50% mayor. Completar esa carga es la primera prioridad de mejora.'),

    ('¿Por qué el CPA de TTPIP es tan alto ($1.106)?',
     'Es propio de la especialización. El Taller de Técnicas y Procesos Innovadores en Pastelería usa ingredientes de alto valor como couverture de chocolate, frutas exóticas y productos de pastelería premium. Es el mismo fenómeno que en un restaurante de alta cocina vs. uno casual: el costo por plato es mayor por la calidad de los insumos.'),

    ('¿Los precios son actuales?',
     'Los precios fueron cargados desde el Excel operacional que usa el equipo de compras del Instituto, correspondiente al período 2026. Se aplicaron conversiones de unidad (GR→kg, CC→L) para estandarizar la comparación. El sistema permite actualizar precios individualmente o mediante carga masiva cuando los proveedores actualicen sus tarifas.'),

    ('¿Esto reemplaza el ERP o el sistema contable?',
     'No. El Sistema de Costos Gastronómicos es una herramienta específica para la gestión operativa de insumos culinarios. No maneja contabilidad, nómina, ni finanzas. Es complementario a cualquier sistema administrativo existente y puede exportar datos para ser integrados con otros sistemas.'),

    ('¿Quién puede usar el sistema? ¿Es seguro?',
     'El sistema tiene autenticación por usuario y contraseña. Actualmente los roles definidos son administrador y consultor. Los datos están almacenados en Supabase (PostgreSQL), que cumple con estándares de seguridad internacionales (SOC 2, ISO 27001). El acceso está protegido con claves de API específicas por proyecto.'),

    ('¿Cuánto cuesta mantener el sistema?',
     'La infraestructura actual tiene un costo mensual estimado menor a USD 50 (plan de Supabase + PythonAnywhere). El desarrollo fue interno, sin licencias externas. Las actualizaciones de datos pueden realizarlas usuarios sin conocimientos de programación, usando la interfaz web o las plantillas de importación Excel que provee el mismo sistema.'),

    ('¿Se puede extender a otras sedes?',
     'Sí, está previsto en el roadmap de largo plazo. La arquitectura permite crear instancias adicionales (una base de datos por sede) o una base de datos compartida con campo de sede. El sistema fue diseñado para ser escalable desde el inicio.'),
]

for pregunta, respuesta in faqs:
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(10)
    p_q.paragraph_format.space_after = Pt(2)
    rq = p_q.add_run('❓  ' + pregunta)
    rq.bold = True; rq.font.name = 'Calibri'; rq.font.size = Pt(11)
    rq.font.color.rgb = BLUE

    p_a = doc.add_paragraph()
    p_a.paragraph_format.left_indent = Inches(0.3)
    p_a.paragraph_format.space_before = Pt(1)
    p_a.paragraph_format.space_after = Pt(6)
    ra = p_a.add_run(respuesta)
    ra.font.name = 'Calibri'; ra.font.size = Pt(11)

    divider()

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# SECCIÓN 4: ADAPTACIÓN POR AUDIENCIA
# ═══════════════════════════════════════════════════════
h1('4.  Adaptación por Tipo de Audiencia')
body('La presentación puede adaptarse según el perfil del público. A continuación, las recomendaciones por audiencia.')

audiencias = [
    ('🎓 Dirección Académica',
     ['Énfasis en slides 6, 7 y 9 (datos, KPIs y análisis por taller).',
      'Destacar el análisis de CPA: permite comparar eficiencia de talleres.',
      'Mencionar los 5 talleres sin datos como brecha de gestión.',
      'Conectar con planificación presupuestaria anual.',
      'Duración recomendada: 20 min + 10 min preguntas.']),
    ('💼 Administración y Finanzas',
     ['Énfasis en slides 6, 7 y 8 (cifras, KPIs y categorías).',
      'Destacar el costo de infraestructura del sistema (< USD 50/mes).',
      'Mostrar la exportación a Excel como herramienta de auditoría.',
      'Conectar con el módulo de presupuesto vs. real (roadmap mediano plazo).',
      'Duración recomendada: 15 min + 5 min preguntas.']),
    ('👨‍🍳 Coordinadores de Taller',
     ['Énfasis en slides 9 y 10 (análisis por taller y top sesiones).',
      'Mostrar demo en vivo: navegación por sesiones, filtros y detalle de recetas.',
      'Explicar cómo se ingresan recetas y se actualizan cantidades.',
      'Destacar la plantilla Excel de importación masiva.',
      'Duración recomendada: 15 min demo + 10 min práctica.']),
    ('🔧 Equipo Técnico / TI',
     ['Énfasis en slide 5 (arquitectura técnica).',
      'Compartir repositorio o código fuente del sistema.',
      'Explicar la estructura de tablas y vistas en Supabase.',
      'Discutir estrategia de backup, monitoreo y escalabilidad.',
      'Duración recomendada: 30 min + 15 min preguntas técnicas.']),
]

for aud, items in audiencias:
    h2(aud, color=GOLD)
    for item in items:
        bullet(item)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# SECCIÓN 5: DATOS RÁPIDOS
# ═══════════════════════════════════════════════════════
h1('5.  Datos Rápidos de Referencia')
body('Para usar durante preguntas o en conversaciones post-presentación.')

datos_rapidos = [
    ('CIFRA', 'DATO', 'CONTEXTO'),
    ('$7.075.872', 'Gasto total mat. prima 2026', '179 OPs · 1.135 sesiones · 9 talleres activos'),
    ('$328', 'CPA promedio por sesión', 'Rango: $254 (TPE) a $1.106 (TTPIP)'),
    ('$1.106', 'CPA más alto (TTPIP)', '40 sesiones · Técnicas Innovadoras Past.'),
    ('$254', 'CPA más bajo (TPE)', '76 sesiones · Panadería de Especialidad'),
    ('$1.178.684', 'Mes más costoso (Mayo 2026)', '16,7% del gasto anual'),
    ('$324.365', 'Mes más bajo (Diciembre 2026)', '4,6% del gasto anual'),
    ('28,8%', 'Participación Abarrote', 'Categoría de mayor gasto: $2.040.961'),
    ('$341.058', 'Ingrediente #1: Mantequilla s/sal', 'Ingrediente de mayor gasto acumulado'),
    ('23.213', 'Registros ingrediente-sesión', 'Nivel de detalle del sistema'),
    ('5 de 14', 'Talleres sin recetas', 'TMA, TAS, TIR, TMI, TEAC · dato pendiente'),
]

t = doc.add_table(rows=len(datos_rapidos), cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.LEFT
for ri, row in enumerate(datos_rapidos):
    for ci, val in enumerate(row):
        cell = t.rows[ri].cells[ci]
        if ri == 0:
            set_cell_bg(cell, '1e1e1e')
        elif ri % 2 == 0:
            set_cell_bg(cell, '141414')
        else:
            set_cell_bg(cell, '1a1a1a')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 2 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.name = 'Calibri'
        r.font.size = Pt(10 if ri > 0 else 11)
        r.bold = (ri == 0 or ci == 0)
        r.font.color.rgb = GOLD if ri == 0 or ci == 0 else (WHITE if ci == 1 else MUTED)

doc.add_paragraph()
body('URL del sistema:  culinary.pythonanywhere.com', muted=False)
body('Archivo PPT:  Culinary_Sistema_Costos_Presentacion.pptx', muted=True)
body('Informe APA:  Culinary_Informe_APA_Sistema_Costos.docx', muted=True)

doc.save(OUT)
print(f'Guia guardada: {OUT}')
