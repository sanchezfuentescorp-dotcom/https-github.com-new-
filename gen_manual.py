# -*- coding: utf-8 -*-
"""Generate Manual de Procedimiento - Sistema de Costos Gastronomicos"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\dmoov\Documents\Programas\costos-gastronomia\Manual_Sistema_Costos_Culinary.docx"

GOLD  = RGBColor(0xFB, 0xBF, 0x24)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLUE  = RGBColor(0x0D, 0x47, 0xA1)
BODY  = RGBColor(0x1F, 0x1F, 0x1F)
H2CLR = RGBColor(0x1A, 0x1A, 0x2E)
H3CLR = RGBColor(0x37, 0x37, 0x37)
NOTEC = RGBColor(0x5D, 0x40, 0x00)
GRAY  = RGBColor(0x88, 0x88, 0x88)

# ── XML helpers ────────────────────────────────────────────────────────────

def cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    # remove existing shd
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    tcPr.append(shd)

def para_shd(para, hex6):
    """Insert shd into pPr in schema-correct position (after pStyle/numPr/keepLines, before tabs/spacing)."""
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn('w:shd')):
        pPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    # Insert before spacing/ind/jc/rPr to keep valid order
    insert_before = (qn('w:tabs'), qn('w:spacing'), qn('w:ind'),
                     qn('w:jc'), qn('w:rPr'))
    ref = None
    for child in pPr:
        if child.tag in insert_before:
            ref = child
            break
    if ref is not None:
        pPr.insert(list(pPr).index(ref), shd)
    else:
        pPr.append(shd)

def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def add_static_toc(doc):
    """Build a static, always-visible index with all sections and subsections."""

    # Top-level sections (title, subsections list)
    entries = [
        ('1. INTRODUCCION', [
            '1.1 Descripcion General',
            '1.2 Arquitectura del Sistema',
            '1.3 Roles de Usuario',
            '1.4 Flujo General del Sistema',
        ]),
        ('2. ACCESO Y CONFIGURACION', [
            '2.1 Primer Acceso - Configuracion Supabase',
            '2.2 Inicio de Sesion',
            '2.3 Persistencia de Modulo Activo',
        ]),
        ('3. NAVEGACION - MENU LATERAL', [
            '3.1 Estructura de Grupos',
            '3.2 Busqueda Global (Ctrl+K)',
        ]),
        ('4. MODULO: DASHBOARD', [
            '4.1 Descripcion',
            '4.2 Filtros Disponibles',
            '4.3 KPIs del Dashboard',
            '4.4 Variacion Mes a Mes',
            '4.5 Graficos del Dashboard',
            '4.6 Exportaciones desde Dashboard',
        ]),
        ('5. MODULO: INGREDIENTES (Gestion Base)', [
            '5.1 Descripcion',
            '5.2 Campos del Ingrediente',
            '5.3 Formula de Precio Bruto',
            '5.4 Alerta de Precio Sin Asignar',
            '5.5 Filtros Disponibles',
            '5.6 Ordenamiento de Columnas',
            '5.7 Historial de Precios por Ingrediente',
            '5.8 Actualizacion Masiva de Precios',
        ]),
        ('6. MODULO: PROVEEDORES (Gestion Base)', [
            '6.1 Descripcion',
            '6.2 Campos del Proveedor',
            '6.3 Historial de Compras por Proveedor',
        ]),
        ('7. MODULO: CATEGORIAS (Gestion Base)', [
            '7.1 Descripcion',
            '7.2 Campos',
        ]),
        ('8. MODULO: CENTROS DE COSTO (Gestion Base)', [
            '8.1 Descripcion',
            '8.2 Campos',
            '8.3 Asignacion',
        ]),
        ('9. MODULO: TIPOS DE TALLER (Operaciones)', [
            '9.1 Descripcion',
            '9.2 Campos',
            '9.3 Metricas Calculadas (v_costo_taller)',
        ]),
        ('10. MODULO: SESIONES DE CLASE (Operaciones)', [
            '10.1 Descripcion',
            '10.2 Campos de la Sesion',
            '10.3 Calculo de Costos - Formulas Exactas',
            '10.4 Ingreso de Ingredientes',
            '10.5 Vista Tabla vs Vista Calendario',
            '10.6 Filtros de Sesiones',
            '10.7 Ordenamiento de Columnas',
            '10.8 Duplicar Sesion',
            '10.9 Plantillas de Sesion',
        ]),
        ('11. MODULO: INVENTARIO (Operaciones)', [
            '11.1 Descripcion',
            '11.2 Campos',
            '11.3 Relacion con Stock Minimo',
        ]),
        ('12. MODULO: LISTA DE COMPRAS (Operaciones)', [
            '12.1 Descripcion',
            '12.2 Fuente de Datos',
            '12.3 Calculo de Cantidad a Comprar',
            '12.4 Estados de un Item en la Lista',
            '12.5 Crear OC desde Lista de Compras',
            '12.6 Exportar a Excel',
        ]),
        ('13. MODULO: CARRERAS / ESPECIALIDADES (Planificacion)', [
            '13.1 Descripcion',
            '13.2 Campos',
            '13.3 Calculo de Costos de Carrera',
            '13.4 Costo Real vs Planificado',
            '13.5 Informe PDF por Carrera',
        ]),
        ('14. MODULO: MODULOS ACADEMICOS (Planificacion)', [
            '14.1 Descripcion',
            '14.2 Campos',
        ]),
        ('15. MODULO: SEMANAS DE CLASES (Planificacion)', [
            '15.1 Descripcion',
            '15.2 Campos',
        ]),
        ('16. MODULO: SIMULADOR DE COSTOS (Analisis)', [
            '16.1 Descripcion',
            '16.2 Modos de Simulacion',
            '16.3 Calculo de Simulacion',
        ]),
        ('17. MODULO: COMPARADOR DE PRECIOS (Analisis)', [
            '17.1 Descripcion',
            '17.2 Visualizacion',
        ]),
        ('18. MODULO: HISTORIAL DE PRECIOS (Analisis)', [
            '18.1 Descripcion',
            '18.2 Visualizacion',
            '18.3 Calculo de Variacion',
        ]),
        ('19. MODULO: PRESUPUESTO VS REAL (Analisis)', [
            '19.1 Descripcion',
            '19.2 Campos del Presupuesto',
            '19.3 Calculo de Ejecucion Presupuestaria',
        ]),
        ('20. MODULO: COMPARATIVA DE PERIODOS (Analisis)', [
            '20.1 Descripcion',
            '20.2 Metricas Comparadas',
        ]),
        ('21. MODULO: ORDENES DE COMPRA (Documentos)', [
            '21.1 Descripcion',
            '21.2 Ciclo de Vida de una OC',
            '21.3 Campos de la OC',
            '21.4 Calculo del Total de la OC',
            '21.5 Registro de Recepcion con Diferencias',
            '21.6 Versiones Historicas de la OC',
            '21.7 Exportaciones de OC',
        ]),
        ('22. MODULO: IMPORTAR EXCEL (Herramientas)', [
            '22.1 Descripcion',
            '22.2 Pestanas de Importacion',
            '22.3 Proceso de Importacion General',
            '22.4 Importacion Masiva de Precios',
        ]),
        ('23. MODULO: BACKUP / EXPORTAR (Herramientas)', [
            '23.1 Descripcion',
            '23.2 Exportaciones Disponibles',
            '23.3 Informe PDF por Carrera',
        ]),
        ('24. MODULO: USUARIOS (Administracion)', [
            '24.1 Descripcion',
            '24.2 Campos',
            '24.3 Seguridad',
        ]),
        ('25. FUNCIONALIDADES TRANSVERSALES', [
            '25.1 Centro de Notificaciones',
            '25.2 Modo Oscuro / Claro',
            '25.3 Cierre Mensual PDF',
            '25.4 Plantillas de Sesion',
        ]),
        ('26. ESTRUCTURA DE DATOS EN SUPABASE', [
            '26.1 Tablas Principales',
            '26.2 Vistas Calculadas (Views)',
        ]),
        ('27. FLUJOS DE TRABAJO COMPLETOS', [
            '27.1 Primera Configuracion del Sistema',
            '27.2 Registro de una Sesion de Clase',
            '27.3 Emision de una Orden de Compra',
            '27.4 Recepcion de Mercaderia',
            '27.5 Actualizacion Masiva de Precios',
        ]),
        ('28. GLOSARIO DE TERMINOS', []),
    ]

    for section_title, subsections in entries:
        # Main section entry
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0)
        # Gold bar on left via left indent + color
        r_num = p.add_run(section_title)
        r_num.font.name = 'Arial'
        r_num.font.size = Pt(11)
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        para_shd(p, 'F0F0F0')

        # Subsection entries
        for sub in subsections:
            ps = doc.add_paragraph()
            ps.paragraph_format.space_before = Pt(0)
            ps.paragraph_format.space_after = Pt(0)
            ps.paragraph_format.left_indent = Cm(1.2)
            rs = ps.add_run(sub)
            rs.font.name = 'Arial'
            rs.font.size = Pt(10)
            rs.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def setup_header_footer(section):
    """Running header (right-aligned) + centered page number footer."""
    hdr = section.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run('Instituto Culinary  |  Sistema de Costos Gastronomicos')
    r.font.name = 'Arial'
    r.font.size = Pt(8.5)
    r.font.color.rgb = GRAY
    r.font.italic = True

    ftr = section.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = fp.add_run('Pagina ')
    r2.font.name = 'Arial'
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    # PAGE field
    r3 = fp.add_run()
    r3.font.name = 'Arial'
    r3.font.size = Pt(9)
    r3.font.color.rgb = GRAY
    b = OxmlElement('w:fldChar')
    b.set(qn('w:fldCharType'), 'begin')
    r3._r.append(b)
    it = OxmlElement('w:instrText')
    it.text = 'PAGE'
    r3._r.append(it)
    e = OxmlElement('w:fldChar')
    e.set(qn('w:fldCharType'), 'end')
    r3._r.append(e)

# ── Content helpers ────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.4)
    para_shd(p, '1A1A1A')
    r = p.add_run(text)
    r.font.color.rgb = GOLD
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.name = 'Arial'
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.color.rgb = H2CLR
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.name = 'Arial'
    return p

def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY
    p.paragraph_format.space_after = Pt(5)
    return p

def formula(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)
    r.font.color.rgb = BLUE
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    para_shd(p, 'EFF6FF')
    return p

def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run('Nota: ' + text)
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = NOTEC
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    para_shd(p, 'FFF8E1')
    return p

def bulleted(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(item)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def flow(doc, steps):
    text = '  >  '.join(steps)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x14, 0x49, 0x6C)
    r.font.bold = True
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    para_shd(p, 'E8F4FD')

def tbl(doc, headers, rows, widths=None):
    nc = len(headers)
    if widths is None:
        total = Inches(6.3)
        widths = [total / nc] * nc
    t = doc.add_table(rows=1, cols=nc)
    t.style = 'Table Grid'
    hr = t.rows[0]
    hr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hr.height = Cm(0.75)
    for i, (c, h) in enumerate(zip(hr.cells, headers)):
        c.width = widths[i]
        cell_bg(c, '1A1A1A')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.color.rgb = WHITE
        r.font.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        dr = t.add_row()
        bg = 'F4F4F4' if ri % 2 == 0 else 'FFFFFF'
        for ci, (c, d) in enumerate(zip(dr.cells, row)):
            c.width = widths[ci]
            cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            r = p.add_run(str(d))
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.font.color.rgb = BODY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ══════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

# A4 page setup
for sec in doc.sections:
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.0)

# Suppress header/footer on cover section
sec0 = doc.sections[0]
sec0.header.is_linked_to_previous = False
sec0.footer.is_linked_to_previous = False
sec0.header.paragraphs[0].clear()
sec0.footer.paragraphs[0].clear()

# ── COVER PAGE ─────────────────────────────────────────────────────────────
def cover_p(doc, text, size, color, bold=True, after=0,
            align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    para_shd(p, '111111')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color
    return p

for _ in range(8):
    cover_p(doc, '', 12, WHITE)

cover_p(doc, 'MANUAL DE PROCEDIMIENTO', 36, GOLD, after=10)
cover_p(doc, 'Sistema de Costos Gastronomicos', 18, WHITE, after=6)
cover_p(doc, '________________________________________', 12, GOLD, bold=False, after=10)
cover_p(doc, 'INSTITUTO CULINARY', 16, WHITE, after=8)
cover_p(doc, 'Version 2.0  -  Junio 2026', 12, RGBColor(0xCC, 0xCC, 0xCC), bold=False, after=4)
cover_p(doc, 'https://culinary.pythonanywhere.com', 10, GOLD, bold=False, after=0)

for _ in range(10):
    cover_p(doc, '', 12, WHITE)

page_break(doc)

# ── CONTENT SECTION (header + footer) ──────────────────────────────────────
new_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
new_sec.page_width = Cm(21)
new_sec.page_height = Cm(29.7)
new_sec.left_margin = Cm(2.54)
new_sec.right_margin = Cm(2.54)
new_sec.top_margin = Cm(2.54)
new_sec.bottom_margin = Cm(2.0)
setup_header_footer(new_sec)

# ── TOC PAGE ───────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tr = tp.add_run('TABLA DE CONTENIDOS')
tr.font.name = 'Arial'
tr.font.size = Pt(16)
tr.font.bold = True
tr.font.color.rgb = H2CLR
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = Pt(16)

add_static_toc(doc)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCION
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '1. INTRODUCCION')

h2(doc, '1.1 Descripcion General')
body(doc, 'El Sistema de Costos Gastronomicos del Instituto Culinary es una aplicacion web de archivo unico (index.html) alojada en PythonAnywhere. Utiliza Supabase como base de datos en la nube mediante su API REST. No requiere instalacion: solo un navegador web moderno.')
body(doc, 'El sistema centraliza el control de costos de materia prima de todas las sesiones de clase, genera ordenes de compra, gestiona inventario, y produce reportes analiticos para la toma de decisiones academicas y financieras.')

h2(doc, '1.2 Arquitectura del Sistema')
bulleted(doc, [
    'Frontend: HTML5 + CSS3 + JavaScript vanilla (archivo unico, sin instalacion ni build)',
    'Base de datos: Supabase (PostgreSQL en la nube, acceso via API REST)',
    'Graficos: Chart.js 4.4',
    'Excel: XLSX 0.18 (lectura y escritura en el navegador)',
    'PDF: jsPDF 2.5 + autotable 3.8',
    'Autenticacion: Sistema propio con hash SHA-256 y sesion por sessionStorage',
    'Almacenamiento local: localStorage para OC, plantillas de sesion, configuracion de navegacion',
])

h2(doc, '1.3 Roles de Usuario')
tbl(doc,
    ['Rol', 'Acceso', 'Descripcion'],
    [
        ['Administrador', 'Total', 'Ve todos los modulos incluyendo Usuarios y configuracion'],
        ['Operador', 'Parcial', 'No ve modulo de Usuarios ni opciones de administracion'],
        ['Consulta', 'Solo lectura', 'Acceso visual sin capacidad de crear, editar ni eliminar'],
    ],
    [Inches(1.5), Inches(1.5), Inches(3.3)]
)

h2(doc, '1.4 Flujo General del Sistema')
flow(doc, ['Crear Ingredientes', 'Asignar Proveedor/Categoria', 'Registrar Sesiones',
           'Costos Automaticos', 'Lista de Compras', 'Emitir OC', 'Dashboard'])

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 2. ACCESO Y CONFIGURACION
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '2. ACCESO Y CONFIGURACION')

h2(doc, '2.1 Primer Acceso - Configuracion Supabase')
body(doc, 'En el primer acceso se muestra la pantalla de configuracion. Se deben ingresar:')
bulleted(doc, [
    'URL del proyecto Supabase (ej: https://xxxx.supabase.co)',
    'API Key (anon key del proyecto)',
])
body(doc, 'Estos datos se guardan en el navegador (localStorage). Si se limpian los datos del navegador, habra que configurar nuevamente.')

h2(doc, '2.2 Inicio de Sesion')
body(doc, 'El sistema solicita usuario y contrasena. Las credenciales se validan contra la tabla "usuarios" en Supabase. La contrasena se hashea con SHA-256 en el navegador antes de comparar: nunca se transmite en texto plano. La sesion persiste durante la ventana del navegador (sessionStorage) y expira al cerrarla.')

h2(doc, '2.3 Persistencia de Modulo Activo')
body(doc, 'Al recargar la pagina, el sistema recuerda el ultimo modulo visitado y lo restaura automaticamente (clave localStorage: culinary_last_sec). El estado expandido/colapsado de cada grupo del menu lateral tambien se preserva (clave: culinary_nav_grps).')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 3. NAVEGACION
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '3. NAVEGACION - MENU LATERAL')

h2(doc, '3.1 Estructura de Grupos')
body(doc, 'El menu lateral esta organizado en grupos colapsables. Hacer clic en el nombre del grupo lo expande o colapsa. El estado de cada grupo se guarda automaticamente en localStorage.')
tbl(doc,
    ['Grupo', 'Modulos que contiene'],
    [
        ['Resumen', 'Dashboard'],
        ['Gestion Base', 'Ingredientes, Proveedores, Categorias, Centros de Costo'],
        ['Operaciones', 'Tipos de Taller, Sesiones de Clase, Inventario, Lista de Compras'],
        ['Planificacion', 'Carreras/Especialidades, Modulos Academicos, Semanas de Clases'],
        ['Analisis', 'Simulador, Comparador, Historial de Precios, Presupuesto vs Real, Comparativa'],
        ['Documentos', 'Ordenes de Compra'],
        ['Herramientas', 'Importar Excel, Backup'],
        ['Administracion', 'Usuarios'],
    ],
    [Inches(1.8), Inches(4.5)]
)

h2(doc, '3.2 Busqueda Global (Ctrl+K)')
body(doc, 'Presionar Ctrl+K (o hacer clic en el icono de lupa) abre el buscador global. Busca en tiempo real en:')
bulleted(doc, [
    'Ingredientes (por nombre o codigo)',
    'Proveedores (por nombre)',
    'Sesiones (por codigo de clase)',
    'Carreras (por nombre)',
    'Ordenes de Compra (por numero de OC o proveedor)',
])
body(doc, 'Al seleccionar un resultado, navega directamente al modulo correspondiente y resalta el registro.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 4. DASHBOARD
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '4. MODULO: DASHBOARD')

h2(doc, '4.1 Descripcion')
body(doc, 'Pantalla de resumen ejecutivo con KPIs, graficos y tablas de resumen. Es el primer modulo que se muestra al iniciar sesion. Todos los datos se actualizan en tiempo real al cambiar los filtros.')

h2(doc, '4.2 Filtros Disponibles')
tbl(doc,
    ['Filtro', 'Descripcion'],
    [
        ['Semana', 'Filtra todas las metricas por una semana academica especifica'],
        ['Tipo de Taller', 'Filtra por tipo de taller (Cocina Caliente, Pasteleria, etc.)'],
        ['Anio', 'Filtra por anio calendario'],
        ['Mes', 'Filtra por mes especifico'],
    ],
    [Inches(1.8), Inches(4.5)]
)

h2(doc, '4.3 KPIs del Dashboard')
tbl(doc,
    ['KPI', 'Calculo', 'Descripcion'],
    [
        ['Ingredientes', 'COUNT(ingrediente)', 'Total de ingredientes en el maestro'],
        ['Proveedores', 'COUNT(proveedor)', 'Total de proveedores activos'],
        ['Tipos de Taller', 'COUNT(taller_tipo)', 'Tipos de taller registrados'],
        ['Sesiones', 'COUNT(v_costo_sesion filtrado)', 'Sesiones en el periodo filtrado'],
        ['Total Alumnos', 'SUM(num_alumnos)', 'Suma de alumnos en el periodo'],
        ['CPA (Costo/Alumno)', 'AVG(costo_por_alumno) donde num_alumnos > 0', 'Promedio del costo de materia prima por alumno'],
        ['Costo Total MP', 'SUM(costo_total)', 'Suma total de materia prima en el periodo'],
        ['Carreras', 'COUNT(v_costo_carrera)', 'Carreras/especialidades registradas'],
        ['Modulos', 'COUNT(modulo)', 'Modulos academicos registrados'],
        ['Alertas Compras', 'COUNT(estado = COMPRAR)', 'Items bajo stock que necesitan reponerse'],
    ],
    [Inches(1.8), Inches(2.2), Inches(2.3)]
)

h2(doc, '4.4 Variacion Mes a Mes')
body(doc, 'Los KPIs de CPA y Costo Total MP muestran una flecha indicadora comparando con el mes anterior inmediato disponible:')
bulleted(doc, [
    'Flecha roja (subida) = el costo aumento respecto al mes anterior',
    'Flecha verde (bajada) = el costo bajo respecto al mes anterior',
])
formula(doc, 'variacion_pct = |((actual - anterior) / anterior) x 100|')

h2(doc, '4.5 Graficos del Dashboard')
tbl(doc,
    ['Grafico', 'Tipo', 'Fuente', 'Descripcion'],
    [
        ['Costo por Tipo de Taller', 'Barras horiz.', 'v_costo_taller', 'CPA promedio por tipo de taller'],
        ['Distribucion por Categoria', 'Dona', 'v_costo_categoria', 'Participacion porcentual de cada categoria en el gasto total'],
        ['Top 10 Ingredientes', 'Barras horiz.', 'v_top_ingredientes', 'Ingredientes con mayor costo acumulado'],
        ['Costos por Carrera', 'Barras apiladas', 'v_costo_carrera', 'Costo de MP y overhead por carrera'],
        ['Evolucion Mensual', 'Lineas doble eje', 'v_costo_sesion', 'CPA mensual (eje izq.) + Total MP mensual (eje der.)'],
        ['Top 5 Sesiones mas caras', 'Tabla', 'v_costo_sesion', 'Sesiones con mayor costo total de materia prima'],
        ['Alertas de Inventario', 'Tabla', 'ingrediente + inventario', 'Ingredientes donde stock actual < stock_minimo'],
        ['Ultimas 5 Sesiones', 'Tabla', 'v_costo_sesion', 'Sesiones mas recientes registradas'],
    ],
    [Inches(1.9), Inches(1.1), Inches(1.3), Inches(2.0)]
)

h2(doc, '4.6 Exportaciones desde Dashboard')
tbl(doc,
    ['Boton', 'Genera', 'Contenido'],
    [
        ['Exportar PDF', 'PDF A4', 'KPIs + tabla de ultimas sesiones'],
        ['Exportar Excel', '.xlsx', 'Datos completos del periodo filtrado'],
        ['Cierre Mensual', 'PDF A4', 'KPIs, detalle de todas las sesiones, Top 5 mas caras y OC del periodo'],
    ],
    [Inches(1.8), Inches(1.3), Inches(3.2)]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 5. INGREDIENTES
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '5. MODULO: INGREDIENTES (Gestion Base)')

h2(doc, '5.1 Descripcion')
body(doc, 'Maestro central de ingredientes. Cada ingrediente es la unidad base del calculo de costos. Toda sesion de clase referencia ingredientes de este maestro.')

h2(doc, '5.2 Campos del Ingrediente')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Codigo', 'Texto', 'Codigo interno (ej: ING001). Usado en importacion masiva'],
        ['Nombre', 'Texto', 'Nombre del ingrediente (ej: Aceite de oliva)'],
        ['Precio Neto', 'Numero', 'Precio sin IVA en CLP. Base de TODOS los calculos de costo'],
        ['Precio Bruto (+IVA)', 'Calculado', 'Precio Neto x 1.19 (IVA 19%). Solo informativo'],
        ['Unidad', 'Texto', 'Unidad de medida (kg, lt, unidad, gramo, etc.)'],
        ['Formato', 'Texto', 'Presentacion del producto (1 kg, 5 lt, caja 12 un., etc.)'],
        ['Categoria', 'Lista', 'Grupo al que pertenece (lacteos, carnes, verduras, etc.)'],
        ['Proveedor', 'Lista', 'Proveedor habitual del ingrediente'],
        ['Sede', 'Lista', 'Santiago o Vina del Mar'],
        ['Stock Minimo', 'Numero', 'Cantidad minima requerida. Activa alerta si el stock actual es inferior'],
    ],
    [Inches(1.8), Inches(1.2), Inches(3.3)]
)

h2(doc, '5.3 Formula de Precio Bruto')
formula(doc, 'Precio Bruto = Precio Neto x 1.19')
note(doc, 'El sistema usa el Precio Neto (sin IVA) para todos los calculos de costo de sesiones. El Precio Bruto se muestra solo como referencia informativa.')

h2(doc, '5.4 Alerta de Precio Sin Asignar')
body(doc, 'El sistema notifica si existen ingredientes con costo_neto = 0 o nulo, ya que afectan la precision de los calculos de costo de sesion.')

h2(doc, '5.5 Filtros Disponibles')
bulleted(doc, ['Busqueda por nombre o codigo', 'Filtro por categoria', 'Filtro por proveedor', 'Filtro por sede'])

h2(doc, '5.6 Ordenamiento de Columnas')
body(doc, 'Hacer clic en los encabezados de columna (Codigo, Ingrediente, Precio Neto, Precio Bruto, Categoria, Proveedor) ordena la tabla de forma ascendente o descendente. La flecha dorada indica la columna y direccion activa.')

h2(doc, '5.7 Historial de Precios por Ingrediente')
body(doc, 'Cada ingrediente lleva un registro historico de precios por mes/anio. Permite ver la evolucion y variacion porcentual entre periodos. Se accede desde el boton de historial en la fila del ingrediente.')

h2(doc, '5.8 Actualizacion Masiva de Precios')
body(doc, 'Desde el modulo Importar Excel > pestana "Actualizar Precios": subir un archivo .xlsx con columnas "codigo" (o "nombre") y "precio_neto". El sistema muestra una vista previa con el precio actual, el precio nuevo y la variacion porcentual antes de confirmar.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 6. PROVEEDORES
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '6. MODULO: PROVEEDORES (Gestion Base)')

h2(doc, '6.1 Descripcion')
body(doc, 'Directorio de proveedores. Se utiliza para asignar ingredientes, generar ordenes de compra y analizar el historial de compras por proveedor.')

h2(doc, '6.2 Campos del Proveedor')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Nombre', 'Texto', 'Razon social o nombre comercial'],
        ['RUT', 'Texto', 'RUT chileno del proveedor (formato: 76.123.456-7)'],
        ['Direccion', 'Texto', 'Direccion fisica del proveedor'],
        ['Correo', 'Email', 'Correo electronico para envio de ordenes de compra'],
        ['Telefono', 'Texto', 'Numero de contacto'],
        ['Sede', 'Lista', 'Santiago o Vina del Mar'],
        ['Condiciones de Pago', 'Texto', 'Ej: 30 dias, contado, 60 dias fecha factura'],
    ],
    [Inches(1.8), Inches(1.2), Inches(3.3)]
)

h2(doc, '6.3 Historial de Compras por Proveedor')
body(doc, 'Desde la vista de proveedor se accede al historial de Ordenes de Compra emitidas, con KPIs de: total de OC, monto total acumulado y promedio por OC.')

# ══════════════════════════════════════════════════════════════════════════
# 7. CATEGORIAS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '7. MODULO: CATEGORIAS (Gestion Base)')

h2(doc, '7.1 Descripcion')
body(doc, 'Grupos para clasificar ingredientes. Se usan en filtros y en el grafico de distribucion de costos por categoria del Dashboard.')

h2(doc, '7.2 Campos')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Nombre', 'Texto', 'Nombre de la categoria (ej: Lacteos, Carnes, Verduras, Especias)'],
        ['Color', 'Color hex', 'Color usado en graficos y badges de identificacion (ej: #10b981)'],
    ],
    [Inches(1.5), Inches(1.3), Inches(3.5)]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 8. CENTROS DE COSTO
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '8. MODULO: CENTROS DE COSTO (Gestion Base)')

h2(doc, '8.1 Descripcion')
body(doc, 'Clasificacion financiera que permite agrupar sesiones de clase por area, proyecto o centro de responsabilidad. Facilita el control presupuestario por unidad organizativa.')

h2(doc, '8.2 Campos')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Nombre', 'Texto', 'Nombre del centro (ej: Cocina Santiago, Pasteleria Vina)'],
        ['Codigo', 'Texto', 'Codigo abreviado para referencia rapida (ej: CC-001)'],
        ['Descripcion', 'Texto', 'Descripcion del area o proyecto al que pertenece'],
    ],
    [Inches(1.5), Inches(1.3), Inches(3.5)]
)

h2(doc, '8.3 Asignacion')
body(doc, 'Los centros de costo se asignan al crear o editar una sesion de clase (campo opcional). La vista de sesiones muestra la columna "Centro de Costo" con el nombre asignado.')
note(doc, 'La vista v_costo_sesion no incluye centro_costo_id. El sistema hace una consulta adicional a la tabla "sesion" para obtener este dato y lo combina en el cliente.')

# ══════════════════════════════════════════════════════════════════════════
# 9. TIPOS DE TALLER
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '9. MODULO: TIPOS DE TALLER (Operaciones)')

h2(doc, '9.1 Descripcion')
body(doc, 'Define los tipos o categorias de taller que se imparten (Cocina Caliente, Pasteleria, Reposteria, Panaderia, etc.). Cada sesion de clase pertenece obligatoriamente a un tipo de taller.')

h2(doc, '9.2 Campos')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Codigo', 'Texto', 'Codigo abreviado usado en el codigo de clase (ej: CCAL, PAST, PAN)'],
        ['Nombre', 'Texto', 'Nombre completo del tipo de taller'],
        ['Color', 'Color hex', 'Color para identificacion visual en graficos y badges'],
    ],
    [Inches(1.5), Inches(1.3), Inches(3.5)]
)

h2(doc, '9.3 Metricas Calculadas (vista v_costo_taller)')
tbl(doc,
    ['Metrica', 'Calculo'],
    [
        ['Numero de Sesiones', 'COUNT(*) de sesiones del tipo'],
        ['Total Alumnos', 'SUM(num_alumnos) de todas las sesiones del tipo'],
        ['Costo Total', 'SUM(costo_total) de todas las sesiones del tipo'],
        ['Costo Promedio/Alumno', 'AVG(costo_por_alumno) de sesiones con num_alumnos > 0'],
    ],
    [Inches(2.3), Inches(4.0)]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 10. SESIONES DE CLASE
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '10. MODULO: SESIONES DE CLASE (Operaciones)')

h2(doc, '10.1 Descripcion')
body(doc, 'Modulo central del sistema. Registra cada clase impartida con sus ingredientes y cantidades utilizadas. El costo se calcula automaticamente en tiempo real conforme se ingresan los datos.')

h2(doc, '10.2 Campos de la Sesion')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion', 'Obligatorio'],
    [
        ['Tipo de Taller', 'Lista', 'Tipo al que pertenece la sesion', 'Si'],
        ['Semana', 'Lista', 'Semana academica de la sesion', 'Si'],
        ['Seccion', 'Numero', 'Numero de seccion (1 a 10)', 'Si'],
        ['N Alumnos', 'Numero', 'Cantidad de alumnos en la clase', 'Si'],
        ['Fecha', 'Fecha', 'Fecha en que se realizo la sesion (YYYY-MM-DD)', 'Si'],
        ['Centro de Costo', 'Lista', 'Centro al que se imputa el costo de la sesion', 'No'],
        ['Codigo Clase', 'Texto', 'Codigo identificador unico (ej: CCAL-01-S01)', 'No'],
        ['Ingredientes', 'Tabla', 'Lista de ingredientes con cantidad por alumno', 'Al menos 1'],
    ],
    [Inches(1.5), Inches(0.9), Inches(2.7), Inches(1.2)]
)

h2(doc, '10.3 Calculo de Costos - Formulas Exactas')
body(doc, 'Estas formulas se aplican en tiempo real en el formulario y en la base de datos mediante la vista v_costo_sesion:')
formula(doc, 'Formula 1 - Cantidad Total por ingrediente:')
formula(doc, '    cantidad_total = cantidad_por_alumno x num_alumnos')
formula(doc, '')
formula(doc, 'Formula 2 - Subtotal de cada ingrediente:')
formula(doc, '    subtotal_ingrediente = cantidad_total x costo_neto_ingrediente')
formula(doc, '')
formula(doc, 'Formula 3 - Costo Total de la Sesion:')
formula(doc, '    costo_total = SUM(subtotal_ingrediente) para todos los ingredientes')
formula(doc, '')
formula(doc, 'Formula 4 - Costo por Alumno:')
formula(doc, '    costo_por_alumno = costo_total / num_alumnos   [solo si num_alumnos > 0]')

h2(doc, '10.4 Ingreso de Ingredientes')
body(doc, 'Por cada ingrediente se ingresa la Cantidad por Alumno (cant_unit). El sistema calcula automaticamente la Cantidad Total = cant_unit x num_alumnos. El costo total de la sesion se actualiza en tiempo real (indicador "Live Cost").')

h2(doc, '10.5 Vista Tabla vs Vista Calendario')
body(doc, 'La vista predeterminada es una tabla ordenable. El boton "Calendario" cambia a una vista de calendario mensual con las sesiones agrupadas por dia. Los botones < > navegan entre meses.')

h2(doc, '10.6 Filtros de Sesiones')
tbl(doc,
    ['Filtro', 'Descripcion'],
    [
        ['Busqueda por codigo', 'Filtra por codigo de clase (busqueda parcial)'],
        ['Tipo de Taller', 'Filtra sesiones por tipo de taller'],
        ['Centro de Costo', 'Filtra sesiones por centro de costo asignado'],
        ['Fecha Desde / Hasta', 'Rango de fechas de la sesion'],
    ],
    [Inches(2.0), Inches(4.3)]
)

h2(doc, '10.7 Ordenamiento de Columnas')
body(doc, 'Hacer clic en los encabezados (Codigo Clase, Taller, Fecha, Alumnos, Costo Total, Costo/Alumno) ordena la tabla alternando ascendente/descendente. La flecha dorada indica la columna activa.')

h2(doc, '10.8 Duplicar Sesion')
body(doc, 'Cada sesion puede duplicarse con el boton de copia. Carga automaticamente todos los ingredientes y cantidades en una nueva sesion editable antes de guardar.')

h2(doc, '10.9 Plantillas de Sesion')
body(doc, 'Permite guardar la lista de ingredientes de una sesion para reutilizarla en futuras sesiones:')
numbered(doc, [
    'Crear o editar una sesion y agregar todos los ingredientes con sus cantidades',
    'Hacer clic en "Guardar como plantilla" y asignar un nombre descriptivo',
    'En futuras sesiones, hacer clic en "Plantillas" para cargar una plantilla guardada',
    'Al cargar, se pre-rellenan todos los ingredientes con sus cantidades originales',
])
note(doc, 'Las plantillas se guardan en el navegador local (localStorage, clave: culinary_ses_templates). Son especificas del dispositivo y navegador: no se sincronizan entre computadores.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 11. INVENTARIO
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '11. MODULO: INVENTARIO (Operaciones)')

h2(doc, '11.1 Descripcion')
body(doc, 'Registro del stock actual de materias primas. Los datos de inventario se usan para calcular la Lista de Compras, comparando lo disponible versus lo necesario para las sesiones programadas.')

h2(doc, '11.2 Campos')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Ingrediente', 'Lista', 'Ingrediente del maestro al que corresponde el stock'],
        ['Cantidad', 'Numero', 'Stock disponible actual (en la unidad del ingrediente)'],
        ['Mes', 'Numero', 'Mes al que corresponde el registro de stock'],
        ['Anio', 'Numero', 'Anio al que corresponde el registro de stock'],
    ],
    [Inches(1.8), Inches(1.2), Inches(3.3)]
)

h2(doc, '11.3 Relacion con Stock Minimo')
body(doc, 'El stock minimo se define en el modulo Ingredientes. El Inventario registra el stock actual. La comparacion entre ambos genera:')
bulleted(doc, [
    'Alertas en el Dashboard (tabla "Alertas de Inventario" y KPI "Alertas Compras")',
    'Estado "COMPRAR" en la Lista de Compras para ingredientes bajo stock',
    'Notificaciones en el panel de campana (icono en la cabecera)',
])

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 12. LISTA DE COMPRAS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '12. MODULO: LISTA DE COMPRAS (Operaciones)')

h2(doc, '12.1 Descripcion')
body(doc, 'Generada automaticamente comparando lo que se necesita para las proximas sesiones programadas versus el stock disponible en inventario. Permite identificar que comprar y en que cantidad.')

h2(doc, '12.2 Fuente de Datos')
body(doc, 'La vista v_lista_compras de Supabase realiza el calculo de requerimientos consolidando sesiones futuras e inventario actual. El sistema lee esta vista para mostrar el estado de cada ingrediente.')

h2(doc, '12.3 Calculo de Cantidad a Comprar')
formula(doc, 'cantidad_necesaria = SUM(cantidad_total) de sesiones futuras del ingrediente')
formula(doc, 'cantidad_a_comprar = MAX(0, cantidad_necesaria - stock_actual)')
formula(doc, 'costo_estimado     = cantidad_a_comprar x costo_neto_ingrediente')

h2(doc, '12.4 Estados de un Item en la Lista')
tbl(doc,
    ['Estado', 'Descripcion', 'Color'],
    [
        ['COMPRAR', 'Cantidad necesaria supera el stock disponible', 'Rojo / Naranja'],
        ['OK', 'Stock suficiente para cubrir todas las necesidades', 'Verde'],
        ['SIN_USO', 'El ingrediente no tiene sesiones futuras programadas', 'Gris'],
    ],
    [Inches(1.5), Inches(3.5), Inches(1.3)]
)

h2(doc, '12.5 Crear OC desde Lista de Compras')
body(doc, 'El boton "Crear OC desde esta lista" toma automaticamente todos los items con estado COMPRAR y pre-rellena el formulario de Nueva Orden de Compra con los productos, cantidades y costos estimados.')

h2(doc, '12.6 Exportar a Excel')
body(doc, 'Genera un archivo .xlsx con la lista de compras filtrada, incluyendo: nombre del ingrediente, unidad, cantidad a comprar, precio unitario neto y costo estimado total por item.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 13. CARRERAS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '13. MODULO: CARRERAS / ESPECIALIDADES (Planificacion)')

h2(doc, '13.1 Descripcion')
body(doc, 'Define las carreras o especialidades del instituto. Calcula el costo total de materia prima por alumno durante toda la carrera, incluyendo un factor de overhead administrativo.')

h2(doc, '13.2 Campos')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Nombre', 'Texto', 'Nombre de la carrera (ej: Chef Profesional, Pasteleria y Panaderia)'],
        ['Semestres', 'Numero', 'Duracion total de la carrera en semestres'],
        ['Overhead (%)', 'Porcentaje', 'Porcentaje adicional sobre el costo de MP por gastos indirectos'],
        ['Modulos', 'Lista mult.', 'Modulos academicos que componen el plan de estudios de la carrera'],
    ],
    [Inches(1.5), Inches(1.0), Inches(3.8)]
)

h2(doc, '13.3 Calculo de Costos de Carrera (vista v_costo_carrera)')
formula(doc, 'costo_materia_prima = SUM(costo_por_alumno de sesiones de los modulos de la carrera)')
formula(doc, 'costo_overhead      = costo_materia_prima x (overhead_porcentaje / 100)')
formula(doc, 'costo_total_alumno  = costo_materia_prima + costo_overhead')

h2(doc, '13.4 Costo Real vs Planificado')
body(doc, 'La columna "Ses. Reales / Costo Real" muestra cuantas sesiones reales se han registrado para la carrera y el costo acumulado real. La cadena de navegacion es: carrera > modulos > tipos de taller > sesiones.')

h2(doc, '13.5 Informe PDF por Carrera')
body(doc, 'El boton de informe (modulo Backup) genera un PDF imprimible con: KPIs de la carrera, estructura de modulos y talleres, costo por modulo, costo total por alumno y costo de overhead.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 14. MODULOS ACADEMICOS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '14. MODULO: MODULOS ACADEMICOS (Planificacion)')

h2(doc, '14.1 Descripcion')
body(doc, 'Agrupa tipos de taller bajo un nombre de modulo curricular. Por ejemplo, "Fundamentos Culinarios" puede incluir los talleres Cocina Caliente + Pasteleria. Los modulos se asignan a carreras para estructurar el plan de estudios.')

h2(doc, '14.2 Campos')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Nombre', 'Texto', 'Nombre del modulo curricular'],
        ['Descripcion', 'Texto', 'Descripcion breve del contenido del modulo'],
        ['Horas', 'Numero', 'Total de horas pedagogicas del modulo'],
        ['Talleres incluidos', 'Lista mult.', 'Tipos de taller que componen este modulo'],
    ],
    [Inches(1.8), Inches(1.2), Inches(3.3)]
)

# ══════════════════════════════════════════════════════════════════════════
# 15. SEMANAS DE CLASES
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '15. MODULO: SEMANAS DE CLASES (Planificacion)')

h2(doc, '15.1 Descripcion')
body(doc, 'Define los periodos semanales del calendario academico. Las sesiones de clase se asocian a una semana especifica, permitiendo filtrar y analizar costos por periodo.')

h2(doc, '15.2 Campos')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Nombre', 'Texto', 'Nombre descriptivo de la semana (ej: Semana 1, Semana 15 - Cierre)'],
        ['Fecha Inicio', 'Fecha', 'Primer dia habil de la semana academica'],
        ['Fecha Fin', 'Fecha', 'Ultimo dia de la semana academica'],
    ],
    [Inches(1.8), Inches(1.2), Inches(3.3)]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 16. SIMULADOR
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '16. MODULO: SIMULADOR DE COSTOS (Analisis)')

h2(doc, '16.1 Descripcion')
body(doc, 'Permite proyectar el impacto financiero de cambios de precio en los costos actuales. Responde la pregunta: "Que pasa con el costo por alumno si los precios de ingredientes aumentan X%?"')

h2(doc, '16.2 Modos de Simulacion')
tbl(doc,
    ['Modo', 'Descripcion'],
    [
        ['Todos los ingredientes', 'Aplica el factor de ajuste uniformemente a todos los ingredientes del sistema'],
        ['Por ingrediente especifico', 'Ajusta unicamente el precio de un ingrediente seleccionado'],
        ['Por categoria', 'Ajusta todos los ingredientes que pertenecen a una categoria especifica'],
    ],
    [Inches(2.3), Inches(4.0)]
)

h2(doc, '16.3 Calculo de Simulacion')
formula(doc, 'factor = 1 + (porcentaje_ajuste / 100)')
formula(doc, 'costo_simulado_taller = costo_actual_taller x factor')
formula(doc, '')
formula(doc, 'Ejemplo: ajuste +15%  =>  factor = 1.15')
formula(doc, '         costo_simulado = costo_actual x 1.15')
body(doc, 'El resultado se presenta en un grafico de barras comparativo: costo actual vs. costo simulado por tipo de taller, identificando cuales son mas sensibles a variaciones de precio.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 17. COMPARADOR
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '17. MODULO: COMPARADOR DE PRECIOS (Analisis)')

h2(doc, '17.1 Descripcion')
body(doc, 'Muestra todos los ingredientes ordenados por precio neto, agrupados por proveedor. Permite identificar que proveedor ofrece el mejor precio para cada tipo de ingrediente y detectar sobrecostos.')

h2(doc, '17.2 Visualizacion')
body(doc, 'Barra de progreso relativa: el ingrediente mas caro del conjunto representa el 100%. Los demas se muestran proporcionalmente, facilitando la comparacion visual inmediata sin necesidad de leer valores exactos.')

# ══════════════════════════════════════════════════════════════════════════
# 18. HISTORIAL DE PRECIOS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '18. MODULO: HISTORIAL DE PRECIOS (Analisis)')

h2(doc, '18.1 Descripcion')
body(doc, 'Registra la evolucion historica de precios de cada ingrediente por mes y anio. Permite detectar tendencias de inflacion, estacionalidades y verificar la consistencia de las actualizaciones de precio.')

h2(doc, '18.2 Visualizacion')
bulleted(doc, [
    'Grafico de linea con el precio neto historico por mes/anio',
    'Tabla con periodo, precio neto y variacion porcentual respecto al periodo anterior',
    'KPIs: precio minimo historico, precio maximo historico, precio promedio historico',
])

h2(doc, '18.3 Calculo de Variacion')
formula(doc, 'variacion_pct = ((precio_actual - precio_anterior) / precio_anterior) x 100')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 19. PRESUPUESTO VS REAL
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '19. MODULO: PRESUPUESTO VS REAL (Analisis)')

h2(doc, '19.1 Descripcion')
body(doc, 'Compara el gasto de materia prima planificado versus el gasto real registrado en sesiones. Permite controlar la ejecucion presupuestaria y tomar medidas correctivas a tiempo.')

h2(doc, '19.2 Campos del Presupuesto')
tbl(doc,
    ['Campo', 'Tipo', 'Descripcion'],
    [
        ['Periodo', 'Texto', 'Mes/Anio o semestre al que aplica el presupuesto (ej: 2026-03)'],
        ['Tipo de Taller', 'Lista', 'Taller al que se asigna el monto presupuestado'],
        ['Monto Presupuestado', 'Numero', 'Gasto planificado en CLP de materia prima'],
    ],
    [Inches(1.8), Inches(1.2), Inches(3.3)]
)

h2(doc, '19.3 Calculo de Ejecucion Presupuestaria')
formula(doc, 'variacion     = monto_real - monto_presupuestado')
formula(doc, 'variacion_pct = (variacion / monto_presupuestado) x 100')
formula(doc, '')
formula(doc, 'variacion > 0  =>  gasto mayor al presupuesto  [ALERTA ROJA]')
formula(doc, 'variacion < 0  =>  gasto menor al presupuesto  [FAVORABLE - VERDE]')
formula(doc, 'variacion = 0  =>  ejecucion exacta al presupuesto')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 20. COMPARATIVA DE PERIODOS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '20. MODULO: COMPARATIVA DE PERIODOS (Analisis)')

h2(doc, '20.1 Descripcion')
body(doc, 'Compara el desempeno de costos entre dos anios diferentes. Muestra KPIs comparativos lado a lado y un grafico de evolucion mensual con los dos periodos superpuestos.')

h2(doc, '20.2 Metricas Comparadas')
tbl(doc,
    ['Metrica', 'Formula de Calculo'],
    [
        ['Total sesiones', 'COUNT(sesiones) por anio'],
        ['Total alumnos', 'SUM(num_alumnos) por anio'],
        ['Costo total MP', 'SUM(costo_total) por anio'],
        ['CPA promedio', 'AVG(costo_por_alumno) por anio'],
        ['Variacion CPA', '((CPA_anio_B - CPA_anio_A) / CPA_anio_A) x 100'],
    ],
    [Inches(2.0), Inches(4.3)]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 21. ORDENES DE COMPRA
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '21. MODULO: ORDENES DE COMPRA (Documentos)')

h2(doc, '21.1 Descripcion')
body(doc, 'Gestion completa del ciclo de vida de las ordenes de compra: desde la creacion hasta la recepcion o anulacion. Los datos se almacenan en el navegador local (localStorage, clave: oc_docs), no en Supabase.')

h2(doc, '21.2 Ciclo de Vida de una OC')
flow(doc, ['BORRADOR', 'EMITIDA', 'RECIBIDA o PARCIAL'])
tbl(doc,
    ['Estado', 'Color', 'Descripcion'],
    [
        ['BORRADOR', 'Gris', 'OC en preparacion, no enviada al proveedor. Editable.'],
        ['EMITIDA', 'Azul', 'OC enviada/comprometida con el proveedor. Se registra version historica.'],
        ['RECIBIDA', 'Verde', 'Mercaderia recibida en su totalidad (100% de los items)'],
        ['PARCIAL', 'Naranja', 'Recepcion incompleta: al menos un item con diferencias'],
        ['ANULADA', 'Rojo', 'OC cancelada. Permanece en el historial pero sin efecto.'],
    ],
    [Inches(1.3), Inches(1.0), Inches(4.0)]
)

h2(doc, '21.3 Campos de la OC')
tbl(doc,
    ['Campo', 'Descripcion'],
    [
        ['N OC', 'Generado automaticamente con formato: OC-YYYYMMDD-HHMMSS'],
        ['Proveedor', 'Seleccionado del maestro de proveedores'],
        ['RUT Proveedor', 'Autocompletado desde el maestro de proveedores'],
        ['Moneda', 'CLP (pesos chilenos), USD o EUR'],
        ['Direccion proveedor', 'Autocompletada desde el maestro de proveedores'],
        ['Fecha de emision', 'Fecha de creacion/emision de la OC'],
        ['Fecha entrega esperada', 'Fecha comprometida por el proveedor para entrega'],
        ['Condiciones de pago', 'Autocompletadas desde el maestro (ej: 30 dias)'],
        ['Notas', 'Observaciones o instrucciones especiales para el proveedor'],
        ['Productos', 'Tabla de items: nombre, unidad, cantidad y precio unitario neto'],
    ],
    [Inches(2.3), Inches(4.0)]
)

h2(doc, '21.4 Calculo del Total de la OC')
formula(doc, 'subtotal_neto = SUM(cantidad_item x precio_unitario_neto) para todos los items')
formula(doc, 'IVA           = subtotal_neto x 0.19   (19%)')
formula(doc, 'total_bruto   = subtotal_neto + IVA')

h2(doc, '21.5 Registro de Recepcion con Diferencias')
body(doc, 'Al recibir una OC en estado EMITIDA, se abre el formulario de recepcion con:')
bulleted(doc, [
    'Tabla item a item: cantidad pedida vs. cantidad recibida (campo editable por item)',
    'Indicador visual por item: completo (100%), parcial (menos del 100%), no recibido (0%)',
    'Porcentaje global de recepcion calculado automaticamente',
    'Estado asignado automaticamente: RECIBIDA si 100%, PARCIAL si menos',
    'Campo de notas para registrar discrepancias, productos danados o faltantes',
])

h2(doc, '21.6 Versiones Historicas de la OC')
body(doc, 'Al emitir una OC se guarda automaticamente una version historica con fecha y estado. Se pueden consultar todas las versiones desde el panel de detalle de la OC, permitiendo auditoria de cambios.')

h2(doc, '21.7 Exportaciones de OC')
tbl(doc,
    ['Funcion', 'Descripcion'],
    [
        ['PDF', 'Genera documento formal: logo, datos del proveedor, tabla de productos, IVA y total'],
        ['Enviar por correo', 'Abre el cliente de correo predeterminado con la OC formateada en el cuerpo'],
        ['Duplicar', 'Crea una nueva OC con los mismos datos y productos (estado BORRADOR)'],
        ['Anular', 'Cambia el estado a ANULADA sin eliminar el registro del historial'],
    ],
    [Inches(1.8), Inches(4.5)]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 22. IMPORTAR EXCEL
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '22. MODULO: IMPORTAR EXCEL (Herramientas)')

h2(doc, '22.1 Descripcion')
body(doc, 'Permite carga masiva de datos desde archivos .xlsx directamente en el navegador. El archivo es procesado localmente por la libreria XLSX, sin enviarse a ningun servidor externo. Los datos validados se envian a Supabase via API REST.')

h2(doc, '22.2 Pestanas de Importacion')
tbl(doc,
    ['Pestana', 'Tabla Destino', 'Columnas Requeridas (*)'],
    [
        ['Ingredientes', 'ingrediente', 'nombre*, unidad*, costo_neto, codigo, categoria, proveedor'],
        ['Sesiones', 'sesion', 'taller_codigo*, semana_nombre*, num_alumnos*, fecha*, seccion'],
        ['Detalle sesion', 'sesion_ingrediente', 'codigo_clase*, codigo_ingrediente*, cantidad_unit*, cantidad_total*'],
        ['Inventario', 'inventario', 'codigo_ingrediente*, cantidad*, mes, anio'],
        ['Actualizar Precios', 'ingrediente (PATCH)', 'codigo (o nombre)*, precio_neto*'],
    ],
    [Inches(1.5), Inches(1.5), Inches(3.3)]
)

h2(doc, '22.3 Proceso de Importacion General')
numbered(doc, [
    'Descargar la plantilla Excel haciendo clic en "Descargar plantilla"',
    'Completar los datos en la plantilla respetando los nombres de columna exactos',
    'Arrastrar el archivo al area de carga o hacer clic para seleccionar el archivo',
    'Revisar la vista previa: filas validas en verde, errores en rojo con descripcion',
    'Corregir los errores en el archivo Excel si los hay y volver a cargar',
    'Hacer clic en "Importar todo a Supabase" para confirmar la importacion',
])

h2(doc, '22.4 Importacion Masiva de Precios (Pestana Avanzada)')
numbered(doc, [
    'Preparar Excel con columnas: "codigo" (o "nombre") y "precio_neto"',
    'Subir el archivo en la pestana "Actualizar Precios"',
    'Revisar la vista previa: precio actual, precio nuevo y variacion % para cada ingrediente',
    'Verificar que todos los ingredientes fueron identificados correctamente',
    'Hacer clic en "Actualizar precios en Supabase" para confirmar',
])
note(doc, 'Solo se actualizan ingredientes donde el precio nuevo difiere del precio actual registrado. Los ingredientes no identificados se marcan en rojo y no se importan.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 23. BACKUP
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '23. MODULO: BACKUP / EXPORTAR (Herramientas)')

h2(doc, '23.1 Descripcion')
body(doc, 'Permite exportar todos los datos del sistema en formato Excel para respaldo externo o analisis en herramientas de terceros. Se recomienda hacer backup periodico de los datos.')

h2(doc, '23.2 Exportaciones Disponibles')
tbl(doc,
    ['Exportacion', 'Contenido', 'Formato'],
    [
        ['Ingredientes', 'Todos los campos del maestro de ingredientes con categoria y proveedor', '.xlsx'],
        ['Sesiones', 'Sesiones con costos calculados (desde la vista v_costo_sesion)', '.xlsx'],
        ['Backup completo', 'Todas las tablas del sistema en hojas separadas de un solo archivo', '.xlsx'],
    ],
    [Inches(1.8), Inches(3.5), Inches(1.0)]
)

h2(doc, '23.3 Informe PDF por Carrera')
body(doc, 'Genera un PDF imprimible con el analisis completo de costos de una carrera seleccionada. Incluye: KPIs de la carrera, estructura de modulos y talleres, costo de materia prima por modulo, costo total por alumno y costo de overhead.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 24. USUARIOS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '24. MODULO: USUARIOS (Administracion)')

h2(doc, '24.1 Descripcion')
body(doc, 'Gestion de cuentas de acceso al sistema. Este modulo es visible exclusivamente para usuarios con rol Administrador.')

h2(doc, '24.2 Campos')
tbl(doc,
    ['Campo', 'Descripcion'],
    [
        ['Nombre', 'Nombre completo visible del usuario en el sistema'],
        ['Usuario', 'Nombre de usuario para iniciar sesion (debe ser unico)'],
        ['Contrasena', 'Se almacena como hash SHA-256: nunca se guarda ni transmite en texto plano'],
        ['Rol', 'Nivel de acceso: "admin", "operador" o "consulta"'],
    ],
    [Inches(2.0), Inches(4.3)]
)

h2(doc, '24.3 Seguridad')
bulleted(doc, [
    'Las contrasenas se hashean con SHA-256 en el navegador antes de enviarse a Supabase',
    'Nunca se almacena ni transmite la contrasena en texto plano',
    'El token de sesion se guarda en sessionStorage y expira al cerrar el navegador',
    'Las politicas de Row Level Security (RLS) de Supabase controlan el acceso a nivel de base de datos',
])

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 25. FUNCIONALIDADES TRANSVERSALES
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '25. FUNCIONALIDADES TRANSVERSALES')

h2(doc, '25.1 Centro de Notificaciones (Campana)')
tbl(doc,
    ['Tipo de Notificacion', 'Condicion que la activa', 'Accion Sugerida'],
    [
        ['Stock bajo', 'stock_actual < stock_minimo del ingrediente', 'Ir a Inventario a actualizar stock'],
        ['OC pendiente', 'OC en estado EMITIDA sin recepcion registrada', 'Ir a Ordenes de Compra'],
        ['Ingredientes sin precio', 'costo_neto = 0 o nulo en el maestro', 'Ir a Ingredientes a asignar precios'],
        ['OC recepcion parcial', 'OC en estado PARCIAL con diferencias', 'Revisar las diferencias registradas'],
    ],
    [Inches(1.9), Inches(2.2), Inches(2.2)]
)

h2(doc, '25.2 Modo Oscuro / Claro')
body(doc, 'El boton de sol/luna en la cabecera alterna entre modo claro y modo oscuro. La preferencia del usuario se guarda en localStorage y se aplica automaticamente al cargar la pagina.')

h2(doc, '25.3 Cierre Mensual PDF')
body(doc, 'Disponible desde el Dashboard con el boton "Cierre Mensual". Genera un PDF A4 con:')
bulleted(doc, [
    'Encabezado negro con titulo dorado y nombre del periodo',
    'KPIs del periodo: sesiones totales, alumnos, costo total MP, CPA promedio',
    'Tabla completa de todas las sesiones del periodo con sus costos',
    'Top 5 sesiones mas caras del periodo',
    'Resumen de Ordenes de Compra emitidas en el periodo',
])
note(doc, 'El periodo del cierre se determina por los filtros activos en el Dashboard (anio + mes). Si no hay filtro de mes activo, el sistema usa el mes mas reciente disponible en los datos.')

h2(doc, '25.4 Plantillas de Sesion')
body(doc, 'Almacenadas en localStorage (clave: culinary_ses_templates). Cada plantilla contiene: nombre asignado, tipo de taller asociado, y lista de ingredientes con sus cantidades por alumno. Son especificas del dispositivo y navegador donde fueron creadas.')

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 26. ESTRUCTURA DE DATOS EN SUPABASE
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '26. ESTRUCTURA DE DATOS EN SUPABASE')

h2(doc, '26.1 Tablas Principales')
tbl(doc,
    ['Tabla', 'Descripcion', 'Campos Clave'],
    [
        ['ingrediente', 'Maestro de ingredientes', 'id, nombre, codigo, costo_neto, unidad, categoria_id, proveedor_id, stock_minimo'],
        ['proveedor', 'Directorio de proveedores', 'id, nombre, rut, correo, condiciones_pago'],
        ['categoria', 'Categorias de ingredientes', 'id, nombre, color'],
        ['centro_costo', 'Centros de costo', 'id, nombre, codigo'],
        ['taller_tipo', 'Tipos de taller', 'id, nombre, codigo, color'],
        ['sesion', 'Sesiones de clase', 'id, taller_tipo_id, semana_id, num_alumnos, fecha, centro_costo_id, cod_clase'],
        ['sesion_ingrediente', 'Detalle de ingredientes por sesion', 'id, sesion_id, ingrediente_id, cantidad_unit, cantidad_total'],
        ['inventario', 'Stock actual', 'id, ingrediente_id, cantidad, mes, anio'],
        ['carrera', 'Carreras/especialidades', 'id, nombre, overhead, duracion_horas'],
        ['modulo', 'Modulos academicos', 'id, nombre, descripcion, horas'],
        ['semana', 'Semanas academicas', 'id, nombre, fecha_inicio, fecha_fin'],
        ['carrera_modulo', 'Relacion N:M carrera-modulo', 'carrera_id, modulo_id'],
        ['modulo_taller', 'Relacion N:M modulo-taller', 'modulo_id, taller_tipo_id'],
        ['presupuesto', 'Presupuestos por periodo', 'id, periodo, taller_tipo_id, monto'],
        ['historial_precios', 'Precios historicos', 'id, ingrediente_id, precio_neto, mes, anio'],
        ['usuarios', 'Usuarios del sistema', 'id, nombre, usuario, password_hash, rol'],
    ],
    [Inches(1.5), Inches(1.8), Inches(3.0)]
)

h2(doc, '26.2 Vistas Calculadas en Supabase (Views)')
tbl(doc,
    ['Vista', 'Descripcion', 'Campos Principales'],
    [
        ['v_costo_sesion', 'Sesiones con costos calculados automaticamente por la BD', 'cod_clase, taller_codigo, fecha, num_alumnos, costo_total, costo_por_alumno'],
        ['v_costo_taller', 'Costo agregado por tipo de taller', 'nombre, num_sesiones, total_alumnos, costo_total, costo_promedio_alumno'],
        ['v_costo_categoria', 'Costo por categoria de ingrediente', 'categoria, costo_total, color'],
        ['v_top_ingredientes', 'Ingredientes ordenados por costo total acumulado', 'nombre, categoria, costo_total'],
        ['v_costo_carrera', 'Costos calculados por carrera con overhead', 'nombre, costo_materia_prima, costo_total, overhead'],
        ['v_lista_compras', 'Necesidades de compra vs stock actual', 'nombre, cantidad_necesaria, stock_actual, cantidad_comprar, costo_neto, estado'],
    ],
    [Inches(1.8), Inches(2.1), Inches(2.4)]
)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 27. FLUJOS DE TRABAJO
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '27. FLUJOS DE TRABAJO COMPLETOS')

h2(doc, '27.1 Flujo: Primera Configuracion del Sistema')
numbered(doc, [
    'Ingresar URL y API Key de Supabase en la pantalla de configuracion inicial',
    'Crear las categorias de ingredientes (Lacteos, Carnes, Verduras, etc.)',
    'Crear los proveedores con sus datos de contacto y condiciones de pago',
    'Crear los ingredientes con precios netos y asignar categoria y proveedor',
    'Crear los tipos de taller (Cocina Caliente, Pasteleria, etc.)',
    'Crear las semanas del calendario academico con fechas de inicio y fin',
    'Crear los modulos academicos y asignarles sus tipos de taller',
    'Crear las carreras/especialidades y asignarles los modulos del plan de estudios',
    'El sistema esta listo para registrar sesiones de clase',
])

h2(doc, '27.2 Flujo: Registro de una Sesion de Clase')
numbered(doc, [
    'Ir al modulo Sesiones de Clase > hacer clic en "Nueva sesion"',
    'Seleccionar el Tipo de Taller y la Semana academica correspondiente',
    'Ingresar el N de Alumnos y la Fecha de la sesion',
    'Opcionalmente: seleccionar Centro de Costo y escribir el Codigo de Clase',
    'Agregar cada ingrediente con su cantidad por alumno',
    'Verificar el costo total calculado en tiempo real en el indicador "Live Cost"',
    'Hacer clic en "Guardar sesion"',
    'La sesion aparece inmediatamente en el Dashboard con todas sus metricas',
])

h2(doc, '27.3 Flujo: Emision de una Orden de Compra')
numbered(doc, [
    'Revisar la Lista de Compras para identificar items con estado "COMPRAR"',
    'Opcionalmente hacer clic en "Crear OC desde esta lista" para pre-rellenar',
    'En el formulario de OC: seleccionar proveedor (RUT, direccion y condiciones se autocompletan)',
    'Revisar y ajustar la lista de productos, cantidades y precios unitarios',
    'Agregar la fecha de entrega esperada y notas si corresponde',
    '"Guardar como borrador" para revision posterior o "Emitir" directamente',
    'Al emitir: el estado cambia a EMITIDA y se guarda una version historica',
    'Descargar PDF de la OC o usar "Enviar por correo" para notificar al proveedor',
])

h2(doc, '27.4 Flujo: Recepcion de Mercaderia')
numbered(doc, [
    'Ir a Ordenes de Compra > seleccionar la OC en estado EMITIDA',
    'Hacer clic en "Registrar Recepcion"',
    'Para cada item de la OC, ingresar la cantidad efectivamente recibida',
    'El sistema calcula el porcentaje de recepcion y asigna estado RECIBIDA o PARCIAL',
    'Agregar notas sobre discrepancias, productos danados o faltantes',
    'Confirmar la recepcion: el estado de la OC se actualiza automaticamente',
    'Actualizar el inventario en el modulo correspondiente si se requiere',
])

h2(doc, '27.5 Flujo: Actualizacion Masiva de Precios')
numbered(doc, [
    'Preparar archivo Excel con columnas: "codigo" (o "nombre") y "precio_neto"',
    'Ir al modulo Importar Excel > pestana "Actualizar Precios"',
    'Arrastrar el archivo al area de carga o hacer clic para seleccionar',
    'Revisar la vista previa: precio actual, precio nuevo y variacion % para cada item',
    'Verificar que todos los ingredientes fueron identificados correctamente',
    'Hacer clic en "Actualizar precios en Supabase" para confirmar',
    'Solo se actualizan los ingredientes con precio diferente al actual',
    'Los costos de todas las sesiones futuras se recalculan automaticamente con los nuevos precios',
])

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 28. GLOSARIO
# ══════════════════════════════════════════════════════════════════════════
h1(doc, '28. GLOSARIO DE TERMINOS')

tbl(doc,
    ['Termino', 'Definicion'],
    [
        ['CPA', 'Costo de Materia Prima por Alumno. Metrica principal del sistema para medir eficiencia de costo.'],
        ['Costo Neto', 'Precio del ingrediente sin IVA. Es la base de todos los calculos de costo del sistema.'],
        ['Costo Bruto', 'Precio del ingrediente con IVA incluido (Neto x 1.19). Solo se muestra como referencia informativa.'],
        ['OC', 'Orden de Compra. Documento formal de solicitud de compra emitido a un proveedor.'],
        ['Sesion', 'Una clase especifica con fecha, numero de alumnos e ingredientes registrados.'],
        ['Taller', 'Tipo de clase impartida (Cocina Caliente, Pasteleria, Reposteria, etc.).'],
        ['v_costo_sesion', 'Vista de Supabase (PostgreSQL) que calcula automaticamente los costos por sesion.'],
        ['Stock Minimo', 'Cantidad umbral definida por ingrediente, bajo la cual el sistema genera una alerta de reposicion.'],
        ['Overhead', 'Porcentaje adicional aplicado al costo de materia prima de una carrera por gastos indirectos.'],
        ['Plantilla', 'Conjunto de ingredientes con cantidades guardado en el navegador para reutilizar en sesiones.'],
        ['Modulo', 'Agrupacion de tipos de taller que forma una unidad curricular del plan de estudios de una carrera.'],
        ['Centro de Costo', 'Unidad organizativa o proyecto al que se imputa el gasto de materia prima de una sesion.'],
        ['localStorage', 'Almacenamiento local del navegador para OC, plantillas y configuracion de navegacion.'],
        ['sessionStorage', 'Almacenamiento de sesion del navegador para el token de autenticacion. Expira al cerrar el browser.'],
        ['Supabase', 'Plataforma de base de datos PostgreSQL en la nube usada como backend del sistema.'],
        ['SHA-256', 'Algoritmo de hash criptografico usado para almacenar y verificar contrasenas de forma segura.'],
        ['IVA', 'Impuesto al Valor Agregado. En Chile corresponde al 19% sobre el precio neto.'],
        ['Live Cost', 'Indicador en tiempo real del costo total acumulado de la sesion mientras se agregan ingredientes.'],
    ],
    [Inches(1.8), Inches(4.5)]
)

# ── SAVE ────────────────────────────────────────────────────────────────────
doc.save(OUT)
print('OK -> ' + OUT)
